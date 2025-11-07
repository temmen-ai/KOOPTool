from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple, Union
from zipfile import ZipFile

from docx import Document
from doc2docx import convert
from lxml import etree as ET

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data containers
# ---------------------------------------------------------------------------


@dataclass
class ParagraphFeatures:
    """Structure that captures every feature extracted from a single paragraph."""

    volgnummer: int
    tekst: str
    indented: int
    indentation: int
    bold: int
    italic: int
    underlined: int
    genummerd: int
    niveau: Union[str, int]
    numId: Union[str, int]
    textparts: List[str] = field(default_factory=list)
    textpartformat_u: str = "Nee"
    textpartformat_b: str = "Nee"
    textpartformat_i: str = "Nee"
    textpartformats_u: List[str] = field(default_factory=list)
    textpartformats_b: List[str] = field(default_factory=list)
    textpartformats_i: List[str] = field(default_factory=list)
    style_name: Optional[str] = None
    alignment: Optional[str] = None
    first_line_indent: int = 0
    hanging_indent: int = 0
    space_before: int = 0
    space_after: int = 0
    tab_stops: List[Dict[str, Union[int, str]]] = field(default_factory=list)
    has_tab_runs: int = 0
    leading_text: str = ""
    leading_run_text: str = ""
    num_properties: Dict[str, Union[str, int, None]] = field(default_factory=dict)

    def to_dict(self) -> dict:
        """Helper so existing pandas-based code can consume the features."""
        return asdict(self)


@dataclass
class TableInfo:
    volgnummer: int
    tekst: str = "TABEL GEVONDEN."


@dataclass
class ImageInfo:
    volgnummer: int
    tekst: str = "AFBEELDING GEVONDEN."
    naam: Optional[str] = None
    rel_id: Optional[str] = None


@dataclass
class DocumentExtraction:
    paragraphs: List[ParagraphFeatures]
    tables: List[TableInfo]
    images: List[ImageInfo]
    source_path: Path


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def ensure_docx(input_path: Union[str, Path], *, remove_original: bool = False) -> Path:
    """
    Make sure the input file is converted to .docx if it is an old .doc.

    Returns the path to the .docx file (which might be the original file if no
    conversion was needed).
    """
    input_path = Path(input_path)
    if input_path.suffix.lower() != ".doc":
        return input_path

    docx_path = input_path.with_suffix(".docx")
    logger.info("Converting %s to %s", input_path, docx_path)
    convert(str(input_path), str(docx_path))
    if not docx_path.exists():
        raise FileNotFoundError(f"Conversie is mislukt: {docx_path} niet gevonden.")

    if remove_original:
        try:
            input_path.unlink()
        except OSError as exc:
            logger.warning("Failed to delete original .doc file %s: %s", input_path, exc)
    return docx_path


def read_document(input_path: Union[str, Path]) -> DocumentExtraction:
    """
    Convert (if needed) and read a Word document, returning all paragraph
    features plus table/image metadata.
    """
    docx_path = ensure_docx(input_path)
    document = Document(docx_path)

    tables, images = _collect_tables_and_images(document)
    numbering_map = _read_numbering_definitions(docx_path)
    paragraphs = _extract_paragraph_features(document, numbering_map)

    return DocumentExtraction(
        paragraphs=paragraphs,
        tables=tables,
        images=images,
        source_path=Path(docx_path),
    )


def paragraphs_to_dicts(paragraphs: Sequence[ParagraphFeatures]) -> List[dict]:
    """Utility fan-out for legacy code that still expects list[dict]."""
    return [paragraph.to_dict() for paragraph in paragraphs]


# ---------------------------------------------------------------------------
# Internals
# ---------------------------------------------------------------------------


def _read_numbering_definitions(docx_path: Path) -> Dict[str, Dict[str, Dict]]:
    """
    Parse word/numbering.xml for num/abstractNum definitions so we can snapshot
    original numbering metadata per numId/level.
    """
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": namespace}
    numbering_map: Dict[str, Dict[str, Dict]] = {}

    try:
        with ZipFile(docx_path) as docx:
            with docx.open("word/numbering.xml") as numbering_xml:
                tree = ET.parse(numbering_xml)
    except (KeyError, FileNotFoundError):
        return numbering_map

    root = tree.getroot()

    abstract_defs: Dict[str, Dict[str, Dict[str, Optional[str]]]] = {}
    for abstract in root.findall("w:abstractNum", ns):
        abstract_id = abstract.get(f"{{{namespace}}}abstractNumId")
        if abstract_id is None:
            continue
        levels: Dict[str, Dict[str, Optional[str]]] = {}
        for lvl in abstract.findall("w:lvl", ns):
            ilvl = lvl.get(f"{{{namespace}}}ilvl")
            if ilvl is None:
                continue
            levels[ilvl] = {
                "numFmt": _get_child_attr(lvl, "w:numFmt", ns),
                "lvlText": _get_child_attr(lvl, "w:lvlText", ns, "val"),
                "start": _get_child_attr(lvl, "w:start", ns),
                "lvlJc": _get_child_attr(lvl, "w:lvlJc", ns),
                "pStyle": _get_child_attr(lvl, "w:pStyle", ns),
            }
        abstract_defs[abstract_id] = levels

    for num in root.findall("w:num", ns):
        num_id = num.get(f"{{{namespace}}}numId")
        if num_id is None:
            continue
        abstract_ref = num.find("w:abstractNumId", ns)
        abstract_id = abstract_ref.get(f"{{{namespace}}}val") if abstract_ref is not None else None

        overrides: Dict[str, Dict[str, Optional[str]]] = {}
        for override in num.findall("w:lvlOverride", ns):
            ilvl = override.get(f"{{{namespace}}}ilvl")
            if ilvl is None:
                continue
            overrides[ilvl] = {
                "startOverride": _get_child_attr(override, "w:startOverride", ns, "val"),
            }

        numbering_map[num_id] = {
            "abstractNumId": abstract_id,
            "levels": abstract_defs.get(abstract_id or "", {}),
            "overrides": overrides,
        }

    return numbering_map


def _collect_tables_and_images(document: Document) -> Tuple[List[TableInfo], List[ImageInfo]]:
    """
    Reproduce the original numbering strategy: maintain a running counter that
    increases for every non-empty paragraph, table or embedded image.
    """
    volgnummer_counter = 0
    tables: List[TableInfo] = []
    images: List[ImageInfo] = []
    paragraph_idx = 0

    for element in document.element.body:
        if element.tag.endswith("tbl"):
            volgnummer_counter += 1
            tables.append(TableInfo(volgnummer=volgnummer_counter))
            continue

        if element.tag.endswith("p"):
            paragraph = document.paragraphs[paragraph_idx]
            if _paragraph_has_image(paragraph):
                volgnummer_counter += 1
                img_name, rel_id = _extract_image_metadata(paragraph)
                images.append(
                    ImageInfo(
                        volgnummer=volgnummer_counter,
                        naam=img_name,
                        rel_id=rel_id,
                    )
                )

            if _paragraph_has_visible_text(paragraph):
                volgnummer_counter += 1

            paragraph_idx += 1

    return tables, images


def _extract_paragraph_features(document: Document, numbering_map: Dict[str, Dict[str, Dict]]) -> List[ParagraphFeatures]:
    """Build the paragraph feature list, closely mirroring the legacy logic."""
    paragraphs: List[ParagraphFeatures] = []
    volgnummer = 0

    for paragraph in document.paragraphs:
        if not _paragraph_has_visible_text(paragraph):
            continue

        volgnummer += 1
        indentation, indented = _extract_indentation(paragraph)
        bold = _bool_to_int(_safe_run_attr(paragraph, "bold"))
        italic = _bool_to_int(_safe_run_attr(paragraph, "italic"))
        underlined = _bool_to_int(_safe_run_attr(paragraph, "underline"))
        genummerd, niveau, num_id = _extract_numbering(paragraph)
        textparts, formats_u, formats_b, formats_i = _extract_textparts(paragraph)
        style_name = getattr(getattr(paragraph, "style", None), "name", None)
        alignment = getattr(paragraph.alignment, "name", None) if paragraph.alignment is not None else None
        first_line_length = paragraph.paragraph_format.first_line_indent
        first_line = _length_to_int(first_line_length)
        hanging = abs(first_line) if first_line < 0 else 0
        space_before = _length_to_int(paragraph.paragraph_format.space_before)
        space_after = _length_to_int(paragraph.paragraph_format.space_after)
        tab_stops = _serialize_tab_stops(paragraph.paragraph_format.tab_stops)
        has_tab_runs = 1 if _paragraph_has_tab_runs(paragraph) else 0
        leading_text = _extract_leading_text(paragraph.text)
        leading_run_text = paragraph.runs[0].text if paragraph.runs else ""
        num_props = _lookup_num_properties(numbering_map, num_id, niveau)

        features = ParagraphFeatures(
            volgnummer=volgnummer,
            tekst=paragraph.text,
            indented=indented,
            indentation=indentation,
            bold=bold,
            italic=italic,
            underlined=underlined,
            genummerd=genummerd,
            niveau=niveau,
            numId=num_id,
            textparts=textparts,
            style_name=style_name,
            alignment=alignment,
            first_line_indent=first_line,
            hanging_indent=hanging,
            space_before=space_before,
            space_after=space_after,
            tab_stops=tab_stops,
            has_tab_runs=has_tab_runs,
            leading_text=leading_text,
            leading_run_text=leading_run_text,
            num_properties=num_props,
        )

        if any(val != "None" for val in formats_u):
            features.textpartformat_u = "Ja"
            features.textpartformats_u = formats_u
        if any(val != "None" for val in formats_b):
            features.textpartformat_b = "Ja"
            features.textpartformats_b = formats_b
        if any(val != "None" for val in formats_i):
            features.textpartformat_i = "Ja"
            features.textpartformats_i = formats_i

        paragraphs.append(features)

    return paragraphs


def _paragraph_has_visible_text(paragraph) -> bool:
    text = paragraph.text.strip(" ")
    return bool(text) and not text.isspace()


def _paragraph_has_image(paragraph) -> bool:
    element = paragraph._element  # noqa: SLF001 (python-docx low level usage)
    drawing = element.xpath(".//*[local-name()='drawing']")
    pict = element.xpath(".//*[local-name()='pict']")
    return bool(drawing or pict)


def _extract_image_metadata(paragraph) -> Tuple[Optional[str], Optional[str]]:
    element = paragraph._element
    doc_props = element.xpath(".//*[local-name()='docPr']")
    img_name = doc_props[0].get("name") if doc_props else None

    blip_nodes = element.xpath(".//*[local-name()='blip']")
    rel_id = None
    if blip_nodes:
        for attr, value in blip_nodes[0].attrib.items():
            if attr.endswith("embed"):
                rel_id = value
                break
    return img_name, rel_id


def _extract_indentation(paragraph) -> Tuple[int, int]:
    left_indent = getattr(paragraph.paragraph_format, "left_indent", None)
    if left_indent is None:
        return 0, 0
    try:
        if left_indent > 0:
            return int(left_indent), 1
    except TypeError:
        pass
    return 0, 0


def _bool_to_int(value: Optional[bool]) -> int:
    return 1 if value else 0


def _safe_run_attr(paragraph, attr: str) -> Optional[bool]:
    try:
        first_run = paragraph.runs[0]
    except IndexError:
        return None

    font = first_run.font
    if attr == "bold":
        return bool(getattr(font, "bold", False))
    if attr == "italic":
        return bool(getattr(font, "italic", False))
    if attr == "underline":
        return bool(getattr(font, "underline", False))
    return None


def _extract_numbering(paragraph) -> Tuple[int, Union[str, int], Union[str, int]]:
    try:
        num_pr = paragraph._element.pPr.numPr  # noqa: SLF001
    except AttributeError:
        num_pr = None

    if num_pr is None:
        return 0, "None", "None"

    niveau = getattr(getattr(num_pr, "ilvl", None), "val", "None")
    num_id = getattr(getattr(num_pr, "numId", None), "val", "None")
    return 1, niveau if niveau is not None else "None", num_id if num_id is not None else "None"


def _extract_textparts(paragraph) -> Tuple[List[str], List[str], List[str], List[str]]:
    if len(paragraph.runs) <= 1:
        return [], [], [], []

    textparts: List[str] = []
    formats_u: List[str] = []
    formats_b: List[str] = []
    formats_i: List[str] = []

    for run in paragraph.runs:
        textparts.append(getattr(run._r, "text", ""))  # noqa: SLF001
        formats_u.append(_run_format_value(run, "u"))
        formats_b.append(_run_format_value(run, "b"))
        formats_i.append(_run_format_value(run, "i"))

    return textparts, formats_u, formats_b, formats_i


def _run_format_value(run, attr: str) -> str:
    try:
        r_pr = run._r.rPr  # noqa: SLF001
        value = getattr(getattr(r_pr, attr), "val")
        return getattr(value, "value", value)
    except AttributeError:
        return "None"


def _get_child_attr(parent, tag: str, ns: Dict[str, str], attr: str = "val") -> Optional[str]:
    child = parent.find(tag, ns)
    if child is None:
        return None
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    attr_key = f"{{{namespace}}}{attr}"
    return child.get(attr_key)


def _length_to_int(length_obj) -> int:
    if length_obj is None:
        return 0
    try:
        return int(length_obj)
    except TypeError:
        try:
            return int(length_obj.pt * 20)  # fallback
        except AttributeError:
            return 0


def _serialize_tab_stops(tab_stops) -> List[Dict[str, Union[int, str]]]:
    serialized: List[Dict[str, Union[int, str]]] = []
    if tab_stops is None:
        return serialized
    for tab in tab_stops:
        try:
            position = int(tab.position)
        except (TypeError, ValueError):
            position = 0
        alignment = getattr(tab.alignment, "name", None)
        serialized.append({"position": position, "alignment": alignment or ""})
    return serialized


def _paragraph_has_tab_runs(paragraph) -> bool:
    return any("\t" in run.text for run in paragraph.runs)


def _extract_leading_text(text: str) -> str:
    stripped = text.lstrip()
    if not stripped:
        return ""
    match = re.match(r"^([\S]+)", stripped)
    return match.group(1) if match else ""


def _lookup_num_properties(numbering_map: Dict[str, Dict[str, Dict]], num_id: Union[str, int], niveau: Union[str, int]) -> Dict[str, Union[str, int, None]]:
    if num_id in (None, "None"):
        return {}
    num_info = numbering_map.get(str(num_id))
    if not num_info:
        return {}

    ilvl = None
    if niveau not in (None, "None"):
        try:
            ilvl = str(int(niveau))
        except (ValueError, TypeError):
            ilvl = str(niveau)

    level_info = num_info.get("levels", {}).get(ilvl or "", {}) if ilvl else {}
    overrides = num_info.get("overrides", {}).get(ilvl or "", {}) if ilvl else {}

    props: Dict[str, Union[str, int, None]] = {
        "abstractNumId": num_info.get("abstractNumId"),
        "numFmt": level_info.get("numFmt"),
        "lvlText": level_info.get("lvlText"),
        "start": level_info.get("start"),
        "lvlJc": level_info.get("lvlJc"),
        "pStyle": level_info.get("pStyle"),
    }
    if overrides.get("startOverride"):
        props["startOverride"] = overrides.get("startOverride")
    return props
