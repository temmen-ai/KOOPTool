#### Versie 2.3, 31-10-2025
#Versie voor tweede pilot bij Daadkracht
#Draait in batch
#
#
import numpy as np
import time
import os
import torch
import sys
from docx import Document
from doc2docx import convert
import pandas as pd
from transformers import BertTokenizer, BertConfig
from DaadkrachtBERT import BERTFineTunerWithFeatures
from openai import OpenAI
from os import path
from shutil import rmtree
from lxml import etree as ET
from zipfile import ZipFile
from tempfile import mkdtemp
import json
import ast
from collections import defaultdict
from pathlib import Path
import re
import shutil
import platform
import ast
import pandas as pd
from typing import List, Dict, Tuple, Optional
from typing import Callable, Optional, List
import pandas as pd
from pathlib import Path
import logging
from datetime import datetime

logger = logging.getLogger(__name__)
bestandsnaam_zonder_ext = ""

ALLOWED_LABELS = [
    "OPTitel",
    "OPAanhef",
    "OPHoofdstukTitel",
    "OPParagraafTitel",
    "OPArtikelTitel",
    "OPLid",
    "StandaardAlinea",
    "Lijstalinea",
    "OPBijlageTitel",
    "OPOndertekening",
    "OPNotaToelichtingTitel",
]

# Strakke system instructie
RULES_TEXT = (
    "Je krijgt zinnen van gemeentelijke documenten.\n"
    "Elke regel heeft een voorgesteld opmaakprofiel.\n"
    "Het fragment bevat N regels (gegeven in het veld 'num_rows').\n"
    "Geef uitsluitend JSON terug met sleutel 'y': een lijst van N profielen,\n"
    "in dezelfde volgorde als de aangeleverde regels.\n"
    "BELANGRIJK: Je geeft altijd exact num_rows profielen terug,\n"
    "en uitsluitend in geldig JSON-formaat."
)

def chatgpt_corrector(sub_df: pd.DataFrame) -> List[str]:
    """Roept het fine-tuned model aan en retourneert correcte profielen."""
    OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]
    client = OpenAI(api_key=OPENAI_API_KEY)
    model_id = "ft:gpt-4.1-mini-2025-04-14:spynk-consulting:profielencorrectie3:CJ2nmEt9"  # <-- jouw FT model
    
    # alleen kolommen gebruiken die in de training zaten
    include_cols = [
        "volgnummer", "tekst", "indented", "numbered",
        "niveau", "numId", "opmaakprofiel", "confidence",
        "fout", "regel"
    ]
    rows = sub_df[include_cols].to_dict(orient="records")

    user_payload = {
        "num_rows": len(rows),
        "rows": rows
    }

    messages = [
        {"role": "system", "content": RULES_TEXT},
        {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
    ]
    resp = client.chat.completions.create(
        model=model_id,
        messages=messages,
        temperature=0.0,
    )
    
    content = resp.choices[0].message.content
    parsed = json.loads(content)
    y = parsed.get("y") or parsed.get("opmaakprofielen")

    if not isinstance(y, list):
        raise RuntimeError(f"Modelantwoord mist sleutel 'y'. Ontvangen: {parsed.keys()}")

    # lengte afdwingen (pad/truncate)
    n = len(rows)
    if len(y) < n:
        y = y + ["StandaardAlinea"] * (n - len(y))
    elif len(y) > n:
        y = y[:n]

    # whitelist afdwingen
    y_fixed = [lab if lab in ALLOWED_LABELS else "StandaardAlinea" for lab in y]
    return y_fixed

def apply_corrections_by_windows(
    df_in: pd.DataFrame,
    *,
    selector: Optional[Callable[[pd.DataFrame], List[dict]]] = None,
    corrector: Optional[Callable[[pd.DataFrame], List[str]]] = None,
    order_col: str = "volgnummer",
    profile_col: str = "opmaakprofiel",
    error_col: str = "fout",
    orig_col: str = "opmaakprofiel_BERT",     # hier bewaren we de originele BERT-voorspelling
    inplace: bool = False,
) -> pd.DataFrame:
    """
    Maakt vensters o.b.v. fouten (via select_chatgpt_vensters) en past correcties toe per venster.
    - Bewaart originele profielen in `orig_col`.
    - Schrijft de (gecorrigeerde) profielen in `profile_col`.
    - Overlappende vensters zijn al samengevoegd door selector (verwacht).
    - Standaard 'corrector' is identity: retourneert de huidige profielen (geen wijziging).
    
    Parameters
    ----------
    df_in : DataFrame met minstens: volgnummer, opmaakprofiel, fout
    selector : functie die vensters bepaalt, signatuur: selector(df) -> List[dict]
               (default: gebruikt global `select_chatgpt_vensters`)
    corrector : functie die één venster-DataFrame -> lijst[str] (gecorrigeerde labels) teruggeeft.
                default: identity (geeft bestaande labels ongewijzigd terug)
    order_col, profile_col, error_col, orig_col : kolomnamen
    inplace : of we df_in in-place aanpassen (True) of een kopie teruggeven (False)

    Returns
    -------
    DataFrame met extra kolom `orig_col` en (eventueel) aangepaste `profile_col`.
    """
    if not inplace:
        df = df_in.copy()
    else:
        df = df_in

    # basis checks
    for c in (order_col, profile_col, error_col):
        if c not in df.columns:
            raise ValueError(f"Verwachte kolom '{c}' ontbreekt.")
    # volgnummer naar numeriek & sorteer stabiel (we behouden index om terug te schrijven)
    df[order_col] = pd.to_numeric(df[order_col], errors="coerce")
    df.sort_values(by=order_col, kind="stable", inplace=True)

    # bewaar originele BERT-profielen
    if orig_col not in df.columns:
        df[orig_col] = df[profile_col]
    else:
        # alleen vullen waar nog leeg
        df[orig_col] = df[orig_col].fillna(df[profile_col])

    # selector bepalen
    if selector is None:
        if "select_chatgpt_vensters" in globals():
            selector = globals()["select_chatgpt_vensters"]
        else:
            raise ValueError("Geen selector opgegeven en 'select_chatgpt_vensters' niet gevonden in globals().")

    # corrector (default: identity)
    if corrector is None:
        def corrector(sub_df: pd.DataFrame) -> List[str]:
            # identity: retourneer huidige profielen ongewijzigd
            return sub_df[profile_col].astype(str).tolist()

    # vensters bepalen
    vensters = selector(df)

    # per venster: subset pakken (volgnummerrange), corrector aanroepen, terugschrijven in volgorde
    for v in vensters:
        start_v = v["start_volgnummer"]
        end_v   = v["end_volgnummer"]
        if start_v is None or end_v is None:
            continue  # defensief

        mask = (df[order_col] >= start_v) & (df[order_col] <= end_v)
        sub = df.loc[mask].copy()
        # sorteren voor deterministische volgorde van de corrector-input
        sub = sub.sort_values(by=order_col, kind="stable")
        idx = sub.index
        corrected = corrector(sub)
        if not isinstance(corrected, list):
            raise RuntimeError("Corrector moet een lijst van labels retourneren.")
        if len(corrected) != len(sub):
            # veilig herstellen: trunc/pad (hier kun je strenger op falen als je wil)
            if len(corrected) > len(sub):
                corrected = corrected[:len(sub)]
            else:
                corrected = corrected + sub[profile_col].astype(str).tolist()[len(corrected):]
        # schrijf gecorrigeerde labels terug op de rijen in dit venster
        df.loc[idx, profile_col] = corrected

    # klaar: df heeft nu orig_col met oude labels en profile_col met (mogelijk) gecorrigeerde labels
    return df

SECTION_TITLES = {"OPHoofdstukTitel", "OPParagraafTitel", "OPArtikelTitel"}
PREV_ANCHORS   = {"OPArtikelTitel", "OPHoofdstukTitel"}
NEXT_ANCHORS   = {
    "OPArtikelTitel", "OPHoofdstukTitel", "OPParagraafTitel",
    "OPOndertekening", "OPBijlageTitel", "OPNotaToelichtingTitel"
}
TERMINALS      = {"OPOndertekening", "OPNotaToelichtingTitel", "OPBijlageTitel"}


def _merge_intervals(intervals: List[Tuple[int,int]]) -> List[Tuple[int,int]]:
    """Merge overlappende/aanliggende intervallen op basis van (start_idx, end_idx) (beide inclusief)."""
    if not intervals:
        return []
    intervals = sorted(intervals, key=lambda x: (x[0], x[1]))
    merged = [intervals[0]]
    for s, e in intervals[1:]:
        last_s, last_e = merged[-1]
        if s <= last_e + 1:  # aanliggend of overlappend
            merged[-1] = (last_s, max(last_e, e))
        else:
            merged.append((s, e))
    return merged


def select_chatgpt_vensters(
    df: pd.DataFrame,
    *,
    order_col: str = "volgnummer",
    profile_col: str = "opmaakprofiel",
    error_col: str = "fout",
    window_after_terminal: int = 5,
) -> List[Dict]:
    """
    Bepaalt per 'fout' (fout==1) de vensters (opeenvolgende volgnummers) die naar ChatGPT gaan,
    volgens regels A, B en C. Alleen C-vensters (na terminal) worden samengevoegd.

    Return:
      List[dict] met o.a.:
        - start_idx, end_idx               (row-indices in gesorteerde df)
        - start_volgnummer, end_volgnummer (grenzen in termen van volgnummer)
        - reason                           ("A", "B-prevnext", "C-merged", ...)
        - covered_error_rows               (lijst row-indices van fouten binnen dit venster)
    """
    if order_col not in df.columns or profile_col not in df.columns or error_col not in df.columns:
        raise ValueError("Vereiste kolommen ontbreken (verwacht: volgnummer, opmaakprofiel, fout).")

    # sorteer op volgnummer (als numeric)
    work = df.copy()
    work[order_col] = pd.to_numeric(work[order_col], errors="coerce")
    work = work.sort_values(by=order_col, kind="stable").reset_index(drop=True)

    n = len(work)
    if n == 0:
        return []

    # hulplijsten
    volg = work[order_col].tolist()
    prof = work[profile_col].astype(str).tolist()
    err  = work[error_col].fillna(0).astype(int).tolist()

    # indices met fout==1
    error_idxs = [i for i, v in enumerate(err) if v == 1]
    if not error_idxs:
        return []

    # eerste sectietitel (voor regel A)
    first_section_idx: Optional[int] = next((i for i, p in enumerate(prof) if p in SECTION_TITLES), None)

    # eerste terminal (voor terminal-fase)
    first_terminal_idx: Optional[int] = next((i for i, p in enumerate(prof) if p in TERMINALS), None)

    intervals: List[Tuple[int,int]] = []
    reasons:  List[str] = []

    # --- Regel A ---
    if first_section_idx is not None:
        if any(i <= first_section_idx for i in error_idxs):
            intervals.append((0, first_section_idx))
            reasons.append("A")
    else:
        if error_idxs:
            intervals.append((0, n - 1))
            reasons.append("A-no-section")

    # --- Regel B en C ---
    for e_idx in error_idxs:
        # Terminal-fase → type C
        if (first_terminal_idx is not None) and (e_idx >= first_terminal_idx):
            start_i = max(0, e_idx - window_after_terminal)
            end_i   = min(n - 1, e_idx + window_after_terminal)
            intervals.append((start_i, end_i))
            reasons.append("B-terminal±5")  # wordt later samengevoegd tot C
            continue

        # Normale B-sectie
        after_first_section = (first_section_idx is not None) and (e_idx > first_section_idx)
        if not after_first_section:
            continue

        prev_i = None
        for i in range(e_idx, -1, -1):
            if prof[i] in PREV_ANCHORS:
                prev_i = i
                break
        if prev_i is None:
            prev_i = 0

        next_i = None
        for i in range(e_idx + 1, n):
            if prof[i] in NEXT_ANCHORS:
                next_i = i
                break
        if next_i is None:
            next_i = n - 1

        intervals.append((prev_i, next_i))
        reasons.append("B-prevnext")

    # --- Resultaten opbouwen ---
    out = []

    # Niet-terminal vensters (A/B) zonder merge
    for (s, e), r in zip(intervals, reasons):
        if r != "B-terminal±5":
            covered_errs = [i for i in error_idxs if s <= i <= e]
            out.append({
                "start_idx": s,
                "end_idx": e,
                "start_volgnummer": int(volg[s]) if pd.notnull(volg[s]) else None,
                "end_volgnummer": int(volg[e]) if pd.notnull(volg[e]) else None,
                "reason": r,
                "covered_error_rows": covered_errs,
            })

    # Terminal-vensters (C) wél mergen
    terminal_intervals = [(s, e) for (s, e), r in zip(intervals, reasons) if r == "B-terminal±5"]
    merged_terminals = _merge_intervals(terminal_intervals)
    for (s, e) in merged_terminals:
        covered_errs = [i for i in error_idxs if s <= i <= e]
        out.append({
            "start_idx": s,
            "end_idx": e,
            "start_volgnummer": int(volg[s]) if pd.notnull(volg[s]) else None,
            "end_volgnummer": int(volg[e]) if pd.notnull(volg[e]) else None,
            "reason": "C-merged",
            "covered_error_rows": covered_errs,
        })

    return out


class DocxHelper:
    def __init__(self, docx_file):
        self.docx_file = docx_file

    #LET OP: deze functie gebruiken we nu niet. 
    def get_xml(self, filename):
        with ZipFile(self.docx_file) as docx:
            with docx.open(filename) as xml:
                # 1. parse into tree
                return ET.parse(xml)

    def set_xml(self, filename, tree, output_filename):
        
        tmp_dir = mkdtemp()

        with ZipFile(self.docx_file) as docx:
            filenames = docx.namelist()
            docx.extractall(tmp_dir)

        tree.write(path.join(tmp_dir, filename), pretty_print=True)

        with ZipFile(output_filename, 'w') as docx:
            for filename in filenames:
                docx.write(path.join(tmp_dir, filename), filename)

        #hiermee wordt die temp dir en alle subdirectories weer verwijderd.
        rmtree(tmp_dir)
        
    #idem als set_xml maar dan schrijft hij document.xml en numbering.xml
    def set_xmls(self, filename, numbering_xml, tree, numbering_xml_tree, output_filename):
        
        tmp_dir = mkdtemp()

        with ZipFile(self.docx_file) as docx:
            filenames = docx.namelist()
            docx.extractall(tmp_dir)
            
        #Toegevoegd omdat anders styles.xml alleen gekopieerd wordt als er opmerkingen zijn
        shutil.copyfile('template/xml/styles.xml', path.join(tmp_dir, 'word/styles.xml'))            

        tree.write(path.join(tmp_dir, filename), pretty_print=True)
        numbering_xml_tree.write(path.join(tmp_dir, numbering_xml), pretty_print=True)

        with ZipFile(output_filename, 'w') as docx:
            for filename in filenames:
                docx.write(path.join(tmp_dir, filename), filename)

        rmtree(tmp_dir)

    #idem als set_xml maar dan schrijft hij document.xml, numbering.xml en comments.xml
    def set_xmls_comments(self, filename, comments_xml, numbering_xml, tree, numbering_xml_tree, output_filename):
        
        tmp_dir = mkdtemp()

        with ZipFile(self.docx_file) as docx:
            filenames = docx.namelist()
            docx.extractall(tmp_dir)

        #Copy comments related files to tmp_dir
        shutil.copyfile(comments_xml, path.join(tmp_dir, 'word/comments.xml'))
        shutil.copyfile('template/xml/commentsExtended.xml', path.join(tmp_dir, 'word/commentsExtended.xml'))
        shutil.copyfile('template/xml/commentsIds.xml', path.join(tmp_dir, 'word/commentsIds.xml'))
        shutil.copyfile('template/xml/styles.xml', path.join(tmp_dir, 'word/styles.xml'))
        shutil.copyfile('template/xml/[Content_Types].xml', path.join(tmp_dir, '[Content_Types].xml'))
        shutil.copyfile('template/xml/document.xml.rels', path.join(tmp_dir, 'word/_rels/document.xml.rels'))

        #append filenames to file list
        filenames.append('word/comments.xml')
        filenames.append('word/commentsExtended.xml')
        filenames.append('word/commentsIds.xml')
        
        tree.write(path.join(tmp_dir, filename), pretty_print=True)
        numbering_xml_tree.write(path.join(tmp_dir, numbering_xml), pretty_print=True)

        with ZipFile(output_filename, 'w') as docx:
            for filename in filenames:
                docx.write(path.join(tmp_dir, filename), filename)

        rmtree(tmp_dir)

def init_model_BERT():
    # Pad naar de opgeslagen modelmap
    #model_dir = "./_BERTmodel2" 
    model_dir = "./finetuned_BERTmodel2_ronde2" 

    # Laad de tokenizer
    tokenizer = BertTokenizer.from_pretrained(model_dir)

    # Laad de BERT-configuratie
    config = BertConfig.from_pretrained(model_dir)

    # Initialiseer je aangepaste model. Zorg dat hidden_size en num_labels overeenkomen met je training
    model = BERTFineTunerWithFeatures(hidden_size=768, num_labels=11, tokenizer=tokenizer) 

    # Laad de opgeslagen modelgewichten
    model.load_state_dict(torch.load(os.path.join(model_dir, "best_model_weights.pth"), map_location=torch.device('cpu')))

    # Verplaats het model naar het juiste apparaat, let op: per OS anders
    systeminfo = platform.uname()
    if systeminfo[0] == 'Windows':
        if torch.cuda.is_available():
            device = torch.device('cuda')
        else:
            device = torch.device('cpu')
    else:
        if not torch.backends.mps.is_available():
            device = torch.device('cpu')
        else:
            device = torch.device("mps")        
    
    model.to(device)
    
    # Zet het model in evaluatiemodus
    model.eval()
    
    return model, tokenizer, device

def get_final_predictions(predictions_per_sentence):
    weight_schemes = {
        5: [0.1, 0.1, 0.1, 0.4, 0.3],
        4: [0.1, 0.1, 0.3, 0.5],  
        3: [0.2, 0.4, 0.4],
        2: [0.5, 0.5],
        1: [1.0]
    }        
        
    final_predictions = {}
    all_predictions = {}

    # Bereken het gewogen gemiddelde voor elke zin
    for sentence_index, prob_list in predictions_per_sentence.items():
        num_predictions = len(prob_list)  # aantal voorspellingen voor deze zin
        weights = np.array(weight_schemes[num_predictions])  

        # initieer met nullen, vorm = (num_classes,)
        weighted_probs = np.zeros_like(prob_list[0][0])  

        individual_preds = []

        for j, (prob, position) in enumerate(prob_list):
            weight = weights[j]
            weighted_probs += prob * weight
            individual_preds.append(np.argmax(prob))

        weighted_probs /= weights.sum()  # normaliseren

        final_pred = int(np.argmax(weighted_probs))
        confidence = float(weighted_probs[final_pred])  # mate van zekerheid

        all_predictions[sentence_index] = individual_preds
        # sla label en confidence samen op
        final_predictions[sentence_index] = (final_pred, confidence)        

    return final_predictions, all_predictions


def predict_document(final_predictions, document_sentences):
    label_mapping = {
        0: "OPTitel",
        1: "OPAanhef",
        2: "OPHoofdstukTitel",
        3: "OPParagraafTitel",
        4: "OPArtikelTitel",
        5: "OPLid",
        6: "StandaardAlinea",
        7: "Lijstalinea",
        8: "OPBijlageTitel",
        9: "OPOndertekening",
        10: "OPNotaToelichtingTitel",
    }
    
    predicted_document = []

    for sentence in document_sentences:
        pred, conf = final_predictions[sentence['volgnummer']]
        opmaakprofiel = label_mapping[pred]

        partformats_u, partformats_b, partformats_i = [], [], []
        if sentence['partformat_u'] == 'Ja':
            partformats_u = sentence['partformats_u']
        if sentence['partformat_b'] == 'Ja':
            partformats_b = sentence['partformats_b']
        if sentence['partformat_i'] == 'Ja':
            partformats_i = sentence['partformats_i']

        predicted_document.append({
            'volgnummer': sentence['volgnummer'],
            'tekst': sentence['tekst'],
#            'tabs': sentence['tabs'],            
            'indented': sentence['indented'],
            'indentation': sentence['indentation'],   
            'bold': sentence['bold'],
            'italic': sentence['italic'],
            'underlined': sentence['underlined'],
            'numbered': sentence['genummerd'],
            'niveau': sentence['niveau'],
            'numId': sentence['numId'],
            'opmaakprofiel': opmaakprofiel,
            'confidence': conf,   # <-- nieuw veld
            'textpartformat_u': sentence['partformat_u'],
            'textpartformat_b': sentence['partformat_b'],
            'textpartformat_i': sentence['partformat_i'],
            'textpartformats_u': partformats_u,
            'textpartformats_b': partformats_b,
            'textpartformats_i': partformats_i,
            'textparts': sentence['textparts']
        })    

    predicted_document_df = pd.DataFrame(predicted_document)
    predicted_document_df.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_1.csv', sep=';', index=False)
    return predicted_document_df


def check_predicted_document(predicted_document_df):
    
    OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]
    client = OpenAI(api_key=OPENAI_API_KEY)

    #Controleer resultaat (nu alleen nog op bestaan OPTitel)
    titelcheck = predicted_document_df.loc[(predicted_document_df['opmaakprofiel'] == 'OPTitel')]

    if len(titelcheck.index) == 0:
        logger.info(f"ChatGPT 4o-mini (finetuned version) heeft de titel bepaald.")
        
        document_text = " ".join(predicted_document_df['tekst'].tolist())

    # Maak de prompt voor het gefinetunede model    
        completion = client.chat.completions.create(
          model="ft:gpt-4o-mini-2024-07-18:spynk-consulting:dktitels:ADCnMiOP",
          messages = [
             {"role": "system", "content": "Het model beantwoordt vragen over gemeentelijke verordeningen."},
             {"role": "user", "content": f"Document tekst: '{document_text}'\nVraag: Wat is de titel van deze verordening?"}
          ]
        )

        # Voeg de nieuwe rij met de vereiste informatie toe aan het begin van de DataFrame
        new_row = pd.DataFrame({
            'volgnummer': [0],
            'tekst': [completion.choices[0].message.content],
            'indented': ['0'],
            'indentation': ['0'],
            'bold': ['0'],
            'italic': ['0'],
            'underlined': ['0'],
            'numbered': ['0'],
            'niveau': ['None'],
            'numId': ['None'],
            'opmaakprofiel': ['OPTitel'],    
            'textpartformat_u': ['Nee'],
            'textpartformat_b': ['Nee'],
            'textpartformat_i': ['Nee'],
            'textpartformats_u': ['[]'],
            'textpartformats_b': ['[]'],
            'textpartformats_i': ['[]'],
            'textparts': ['[]']
        })


        # Combineer de nieuwe rij met het bestaande DataFrame
        checked_document_df = pd.concat([new_row, predicted_document_df], ignore_index=True)
                
        checked_document_df.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_2.csv', sep=';', index=False)
        
    else:
        checked_document_df = predicted_document_df
        checked_document_df.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_2.csv', sep=';', index=False)
        
    return checked_document_df

def add_table_and_image_info(temp_df, tables_info, images_info):
    # Combineer beide lijsten en sorteer op volgnummer
    combined_info = tables_info + images_info
    combined_info = sorted(combined_info, key=lambda x: x['volgnummer'])
    
    # Maak een kopie van de oorspronkelijke dataframe
    updated_document_df = temp_df.copy()

    # Voeg elke info (tabel of afbeelding) toe op basis van volgnummer
    for info in combined_info:
        volgnummer = info['volgnummer']
        tekst = info['tekst']

        # Basisrij opbouwen
        new_row_data = {
            'volgnummer': [volgnummer],
            'tekst': [tekst],
            'indented': [0],
            'bold': [0],
            'italic': [0],
            'underlined': [0],
            'numbered': [0],
            'niveau': ['None'],
            'numId': ['None'],
            'opmaakprofiel': ['StandaardAlinea'],
            'textpartformat_u': ['Nee'],
            'textpartformat_b': ['Nee'],
            'textpartformat_i': ['Nee'],
            'textpartformats_u': ['[]'],
            'textpartformats_b': ['[]'],
            'textpartformats_i': ['[]'],
            'textparts': ['[]'],
            'fout': [2],
            'opsommingstype': ['']
        }

        # Regeltekst afhankelijk van type
        if 'naam' in info:  # afbeelding
            new_row_data['regel'] = [f"Kopieer afbeelding uit inputbestand ({info.get('naam', '')})."]
        else:  # tabel
            new_row_data['regel'] = ["Kopieer tabel uit inputbestand en maak eventueel opnieuw op."]

        # Maak dataframe van deze nieuwe rij
        new_row = pd.DataFrame(new_row_data)

        # Voeg de nieuwe rij toe op de exacte positie in de dataframe
        updated_document_df = pd.concat(
            [
                updated_document_df[updated_document_df['volgnummer'] < volgnummer],
                new_row,
                updated_document_df[updated_document_df['volgnummer'] >= volgnummer].assign(
                    volgnummer=lambda df: df['volgnummer'] + 1
                )
            ]
        ).reset_index(drop=True)

    # Optioneel: sla het resultaat op naar een CSV-bestand
    updated_document_df.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_3.csv', sep=';', index=False)

    return updated_document_df


def check_structuur_recheck(checked_document_df: pd.DataFrame,
                            *,
                            volg_col: str = "volgnummer",
                            profile_col: str = "opmaakprofiel",
                            fout_col: str = "fout",
                            regel_col: str = "regel",
                            write_csv_path: str = None) -> pd.DataFrame:
    """
    Her-check van de documentstructuur:
      - reset fout==1 -> 0 (fout==2 blijft staan),
      - reset 'regel' (laat tekst eventueel staan bij fout==2),
      - voer alle regels opnieuw uit (zoals in je bestaande implementatie),
      - retourneert het geannoteerde DataFrame (en schrijft optioneel weg).
    """
    # Kopie en nette index
    df = checked_document_df.copy()
    df = df.reset_index(drop=True)

    # Basis-kolommen garanderen
    if fout_col not in df.columns:
        df[fout_col] = 0
    if regel_col not in df.columns:
        df[regel_col] = ""

    # Reset fouten en regels (2 blijft staan)
    mask_fout2 = (df[fout_col] == 2)
    df.loc[~mask_fout2 & (df[fout_col] == 1), fout_col] = 0
    df.loc[~mask_fout2, regel_col] = ""  # 'regel' legen waar geen tafel/afbeelding

    # 'numbered' defensief (sommige CSV's hebben die kolom niet)
    if "numbered" not in df.columns:
        df["numbered"] = 0

    # Eventueel volgnummer numeriek en sorteren (stabiel)
    if volg_col in df.columns:
        df[volg_col] = pd.to_numeric(df[volg_col], errors="coerce")
        df.sort_values(by=volg_col, kind="stable", inplace=True)
        df.reset_index(drop=True, inplace=True)

    # Regels en allowed transitions (zoals in je code)
    allowed_transitions = {
        'OPTitel': ['OPAanhef'],
        'OPAanhef': ['OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel'],
        'OPHoofdstukTitel': ['OPParagraafTitel', 'OPArtikelTitel', 'OPLid', 'StandaardAlinea', 'Lijstalinea'],
        'OPParagraafTitel': ['OPArtikelTitel'],
        'OPArtikelTitel': ['OPLid', 'StandaardAlinea', 'Lijstalinea'],
        'OPLid': ['OPLid', 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel',
                  'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'StandaardAlinea': ['StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel',
                            'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'Lijstalinea': ['StandaardAlinea', 'Lijstalinea', 'OPLid', 'OPHoofdstukTitel', 'OPParagraafTitel',
                        'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'OPOndertekening': ['OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'OPBijlageTitel': ['StandaardAlinea', 'Lijstalinea'],
        'OPNotaToelichtingTitel': ['StandaardAlinea', 'Lijstalinea']
    }

    regels = {
        1: "Het eerste profiel moet 'OPTitel' zijn.",
        2: "'OPTitel' mag maar één keer voorkomen.",
        3: "Na 'OPTitel' komt altijd 'OPAanhef'.",
        4: "Na 'OPAanhef' kan 'OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel' of 'OPArtikelTitel' volgen.",
        5: "Na 'OPHoofdstukTitel' kan 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid', 'StandaardAlinea' of 'Lijstalinea' volgen.",
        6: "Na 'OPParagraafTitel' kan 'OPArtikelTitel' volgen.",
        7: "Na 'OPArtikelTitel' kan 'OPLid', 'StandaardAlinea' of 'Lijstalinea' volgen.",
        8: "Na 'OPLid' kunnen 'OPLid', 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        9: "Na 'StandaardAlinea' kunnen 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        10: "Na 'Lijstalinea' kunnen 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        11: "Na 'OPOndertekening' kunnen 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        12: "Na 'OPBijlageTitel' kunnen 'StandaardAlinea' of 'Lijstalinea' volgen.",
        13: "Als 'OPBijlageTitel' is geweest, dan mag daarna niet meer voorkomen: 'OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel' of 'OPLid' voorkomen.",
        14: "Na 'OPNotaToelichtingTitel' kunnen 'StandaardAlinea' of 'Lijstalinea' volgen.",
        15: "Als 'OPNotaToelichtingTitel' is geweest, dan mag daarna niet meer voorkomen: 'OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid' of 'OPBijlageTitel' voorkomen.",
        16: "Het profiel 'OPNotaToelichtingTitel' mag maar één keer voorkomen.",
        17: "StandaardAlinea kan geen opsomming zijn."
    }

    # Trackers
    optitel_seen = False
    opnotatoelichting_seen = False    # Regel 16
    opbijlage_seen = False            # Regel 13
    opnotatoelichting_geweest = False # Regel 15

    # Check per rij (skip fout==2)
    for i in range(len(df)):
        current_profile = df.loc[i, profile_col]

        # Skip tabellen/afbeeldingen
        if df.loc[i, fout_col] == 2:
            continue

        # Regel 1
        if i == 0 and current_profile != 'OPTitel':
            df.loc[i, fout_col] = 1
            df.loc[i, regel_col] = f"1: {regels[1]}"
            continue

        # Regel 2
        if current_profile == 'OPTitel':
            if optitel_seen:
                df.loc[i, fout_col] = 1
                df.loc[i, regel_col] = f"2: {regels[2]}"
                continue
            optitel_seen = True

        # Regel 3
        if i > 0 and df.loc[i-1, profile_col] == 'OPTitel' and current_profile != 'OPAanhef':
            # Automatische correctie als volgende rij wel OPAanhef is
            if i + 1 < len(df) and df.loc[i+1, profile_col] == 'OPAanhef':
                df.loc[i, profile_col] = 'OPAanhef'
                df.loc[i, regel_col] = "Automatisch aangepast naar OPAanhef (regel 3)."
                df.loc[i, fout_col] = 0
            else:
                df.loc[i, fout_col] = 1
                df.loc[i, regel_col] = f"3: {regels[3]}"
            continue

        # Allowed transitions & regel 17
        if i > 0:
            previous_profile = df.loc[i-1, profile_col]

            # Regel 17
            if current_profile == 'StandaardAlinea' and int(df.loc[i, 'numbered']) == 1:
                df.loc[i, fout_col] = 1
                df.loc[i, regel_col] = f"17: {regels[17]}"

            # Allowed transitions
            if previous_profile in allowed_transitions:
                if current_profile not in allowed_transitions[previous_profile]:
                    if previous_profile == 'OPAanhef':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"4: {regels[4]}"
                    elif previous_profile == 'OPHoofdstukTitel':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"5: {regels[5]}"
                    elif previous_profile == 'OPParagraafTitel':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"6: {regels[6]}"
                    elif previous_profile == 'OPArtikelTitel':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"7: {regels[7]}"
                    elif previous_profile == 'OPLid':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"8: {regels[8]}"
                    elif previous_profile == 'StandaardAlinea':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"9: {regels[9]}"
                    elif previous_profile == 'Lijstalinea':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"10: {regels[10]}"
                    elif previous_profile == 'OPOndertekening':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"11: {regels[11]}"
                    elif previous_profile == 'OPBijlageTitel':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"12: {regels[12]}"
                    elif previous_profile == 'OPNotaToelichtingTitel':
                        df.loc[i, fout_col] = 1
                        df.loc[i, regel_col] = f"14: {regels[14]}"
                    continue

        # Regel 13: na bijlage geen titels/leden
        if opbijlage_seen and current_profile in ['OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid']:
            df.loc[i, fout_col] = 1
            df.loc[i, regel_col] = f"13: {regels[13]}"
            continue
        if current_profile == 'OPBijlageTitel':
            opbijlage_seen = True

        # Regel 15: na nota geen eerdere types/OPBijlageTitel
        if opnotatoelichting_geweest and current_profile in ['OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid', 'OPBijlageTitel']:
            df.loc[i, fout_col] = 1
            df.loc[i, regel_col] = f"15: {regels[15]}"
            continue
        if current_profile == 'OPNotaToelichtingTitel':
            opnotatoelichting_geweest = True

        # Regel 16: OPNotaToelichtingTitel max 1x
        if current_profile == 'OPNotaToelichtingTitel':
            if opnotatoelichting_seen:
                df.loc[i, fout_col] = 1
                df.loc[i, regel_col] = f"16: {regels[16]}"
                continue
            opnotatoelichting_seen = True

    if write_csv_path:
        df.to_csv(write_csv_path, sep=';', index=False)

    return df


def check_structuur(checked_document_df, tables_info, images_info):
    # Maak een kopie van de input dataframe en voeg de kolommen 'fout' en 'regel' toe
    temp_df = checked_document_df.copy()
    temp_df['fout'] = 0
    temp_df['regel'] = ''

    # Voeg tabellen en afbeeldingen in
    annotated_document_df = add_table_and_image_info(temp_df, tables_info, images_info)
    
    # Definieer de regels die de volgorde bepalen
    allowed_transitions = {
        'OPTitel': ['OPAanhef'],
        'OPAanhef': ['OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel'],
        'OPHoofdstukTitel': ['OPParagraafTitel', 'OPArtikelTitel', 'OPLid', 'StandaardAlinea', 'Lijstalinea'],
        'OPParagraafTitel': ['OPArtikelTitel'],
        'OPArtikelTitel': ['OPLid', 'StandaardAlinea', 'Lijstalinea'],
        'OPLid': ['OPLid', 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel',
                  'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'StandaardAlinea': ['StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel',
                            'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'Lijstalinea': ['StandaardAlinea', 'Lijstalinea', 'OPLid', 'OPHoofdstukTitel', 'OPParagraafTitel',
                        'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'OPOndertekening': ['OPOndertekening', 'OPBijlageTitel', 'OPNotaToelichtingTitel'],
        'OPBijlageTitel': ['StandaardAlinea', 'Lijstalinea'],
        'OPNotaToelichtingTitel': ['StandaardAlinea', 'Lijstalinea']
    }

    # Definieer de regels met hun beschrijvingen
    regels = {
        1: "Het eerste profiel moet 'OPTitel' zijn.",
        2: "'OPTitel' mag maar één keer voorkomen.",
        3: "Na 'OPTitel' komt altijd 'OPAanhef'.",
        4: "Na 'OPAanhef' kan 'OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel' of 'OPArtikelTitel' volgen.",
        5: "Na 'OPHoofdstukTitel' kan 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid', 'StandaardAlinea' of 'Lijstalinea' volgen.",
        6: "Na 'OPParagraafTitel' kan 'OPArtikelTitel' volgen.",
        7: "Na 'OPArtikelTitel' kan 'OPLid', 'StandaardAlinea' of 'Lijstalinea' volgen.",
        8: "Na 'OPLid' kunnen 'OPLid', 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        9: "Na 'StandaardAlinea' kunnen 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        10: "Na 'Lijstalinea' kunnen 'StandaardAlinea', 'Lijstalinea', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        11: "Na 'OPOndertekening' kunnen 'OPOndertekening', 'OPBijlageTitel' of 'OPNotaToelichtingTitel' volgen.",
        12: "Na 'OPBijlageTitel' kunnen 'StandaardAlinea' of 'Lijstalinea' volgen.",
        13: "Als 'OPBijlageTitel' is geweest, dan mag daarna niet meer voorkomen: 'OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel' of 'OPLid' voorkomen.",
        14: "Na 'OPNotaToelichtingTitel' kunnen 'StandaardAlinea' of 'Lijstalinea' volgen.",
        15: "Als 'OPNotaToelichtingTitel' is geweest, dan mag daarna niet meer voorkomen: 'OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid' of 'OPBijlageTitel' voorkomen.",
        16: "Het profiel 'OPNotaToelichtingTitel' mag maar één keer voorkomen.",
        17: "StandaardAlinea kan geen opsomming zijn."
    }
    
    # Trackers voor regels die maar één keer mogen voorkomen
    optitel_seen = False
    opnotatoelichting_seen = False  # Voor regel 16
    opbijlage_seen = False  # Voor regel 13
    opnotatoelichting_geweest = False  # Voor regel 15

    # Doorloop elke rij in het dataframe en check de volgorde
    for i in range(len(annotated_document_df)):
        current_profile = annotated_document_df.loc[i, 'opmaakprofiel']
        
        # Sla tabellen/afbeeldingen over
        if annotated_document_df.loc[i, 'fout'] == 2:
            continue
        
        # Regel 1: Het eerste profiel moet "OPTitel" zijn
        if i == 0 and current_profile != 'OPTitel':
            annotated_document_df.loc[i, 'fout'] = 1
            annotated_document_df.loc[i, 'regel'] = f"1: {regels[1]}"
            continue
        
        # Regel 2: "OPTitel" mag maar één keer voorkomen
        if current_profile == 'OPTitel':
            if optitel_seen:
                annotated_document_df.loc[i, 'fout'] = 1
                annotated_document_df.loc[i, 'regel'] = f"2: {regels[2]}"
                continue
            optitel_seen = True
        
        # Regel 3: Na "OPTitel" komt altijd "OPAanhef"
        if i > 0 and annotated_document_df.loc[i-1, 'opmaakprofiel'] == 'OPTitel' and current_profile != 'OPAanhef':
            # Kijk of de volgende rij wel OPAanhef is
            if i + 1 < len(annotated_document_df) and annotated_document_df.loc[i+1, 'opmaakprofiel'] == 'OPAanhef':
                # Corrigeer deze rij in plaats van fout melden
                annotated_document_df.loc[i, 'opmaakprofiel'] = 'OPAanhef'
                annotated_document_df.loc[i, 'regel'] = "Automatisch aangepast naar OPAanhef (regel 3)."
                annotated_document_df.loc[i, 'fout'] = 0
            else:
                # Dan wel foutmelding
                annotated_document_df.loc[i, 'fout'] = 1
                annotated_document_df.loc[i, 'regel'] = f"3: {regels[3]}"
            continue
        
        # Controleer andere regels op volgorde
        if i > 0:
            previous_profile = annotated_document_df.loc[i-1, 'opmaakprofiel']
            
            # Regel 17: StandaardAlinea kan geen opsomming zijn
            if current_profile == 'StandaardAlinea' and annotated_document_df.loc[i, 'numbered'] == 1:
                annotated_document_df.loc[i, 'fout'] = 1
                annotated_document_df.loc[i, 'regel'] = f"17: {regels[17]}"
            
            # Controleer allowed_transitions
            if previous_profile in allowed_transitions:
                if current_profile not in allowed_transitions[previous_profile]:
                    if previous_profile == 'OPAanhef':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"4: {regels[4]}"
                    elif previous_profile == 'OPHoofdstukTitel':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"5: {regels[5]}"
                    elif previous_profile == 'OPParagraafTitel':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"6: {regels[6]}"
                    elif previous_profile == 'OPArtikelTitel':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"7: {regels[7]}"
                    elif previous_profile == 'OPLid':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"8: {regels[8]}"
                    elif previous_profile == 'StandaardAlinea':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"9: {regels[9]}"
                    elif previous_profile == 'Lijstalinea':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"10: {regels[10]}"
                    elif previous_profile == 'OPOndertekening':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"11: {regels[11]}"
                    elif previous_profile == 'OPBijlageTitel':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"12: {regels[12]}"
                    elif previous_profile == 'OPNotaToelichtingTitel':
                        annotated_document_df.loc[i, 'fout'] = 1
                        annotated_document_df.loc[i, 'regel'] = f"14: {regels[14]}"
                    continue
        
        # Regel 13: OPBijlageTitel blokkeert bepaalde profielen erna
        if opbijlage_seen and current_profile in ['OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid']:
            annotated_document_df.loc[i, 'fout'] = 1
            annotated_document_df.loc[i, 'regel'] = f"13: {regels[13]}"
            continue
        if current_profile == 'OPBijlageTitel':
            opbijlage_seen = True
        
        # Regel 15: OPNotaToelichtingTitel blokkeert bepaalde profielen erna
        if opnotatoelichting_geweest and current_profile in ['OPAanhef', 'OPHoofdstukTitel', 'OPParagraafTitel', 'OPArtikelTitel', 'OPLid', 'OPBijlageTitel']:
            annotated_document_df.loc[i, 'fout'] = 1
            annotated_document_df.loc[i, 'regel'] = f"15: {regels[15]}"
            continue
        if current_profile == 'OPNotaToelichtingTitel':
            opnotatoelichting_geweest = True
        
        # Regel 16: OPNotaToelichtingTitel mag maar één keer
        if current_profile == 'OPNotaToelichtingTitel':
            if opnotatoelichting_seen:
                annotated_document_df.loc[i, 'fout'] = 1
                annotated_document_df.loc[i, 'regel'] = f"16: {regels[16]}"
                continue
            opnotatoelichting_seen = True
            
    annotated_document_df.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_4.csv', sep=';', index=False)
    logger.info("Structuur document gecontroleerd.")
    return annotated_document_df



def post_numbering_cleanup(df: pd.DataFrame) -> pd.DataFrame:
    """
    NA nummeringslogica.

    1) Afgebroken regel samenvoegen (tussen twee gelijke numIds)
    2) Valse nieuwe lijst (zelfde profiel) via numId/bron_numId-sprong, alleen als type gelijk blijft
    3) Collapse runs: OPLid/Lijstalinea zonder numId/numbered → plak bij laatste anker (+ textparts/format merge)
    4) Lokale gelijktrekking: zelfde opsommingstype + andere bron_numId → niveau = vorige
    5) Indent fine-tune: zelfde profiel/type + beide numbered
       - indented stijgt → niveau = vorige + 1
       - indented gelijk → niveau = vorige
       - daling → niet verlagen
    6) (NIEUW) Hiërarchie op basis van (opsommingstype, bron_numId) als indented geen signaal geeft:
       - eerste regel blok: niveau behouden
       - nieuw (type, bronId) in blok: niveau = vorige + 1
       - zelfde (type, bronId): niveau gelijk
       - heroptreden van eerder (type, bronId): spring terug naar niveau van die eerdere

    Aanvulling: Als in een contigu blok (zelfde profiel ∈ {OPLid, Lijstalinea} en numbered==1)
    het hele blok hetzelfde bron_numId én hetzelfde opsommingstype heeft, dan negeren we
    indented volledig in Case 2 en Case 5.
    """
    PROFILES = {"OPLid", "Lijstalinea"}
    def _safe_int(v, default=0):
        try:
            x = pd.to_numeric([v], errors="coerce")[0]
            if pd.isna(x):
                return default
            return int(x)
        except Exception:
            return default

    def _valid_numid(x) -> bool:
        if pd.isna(x):
            return False
        s = str(x).strip().lower()
        return s not in {"", "none", "nan"}

    def _norm_type(x):
        return str(x or "").strip().lower()

    def _valid_bron(s):
        s = str(s or "").strip().lower()
        return s not in {"", "none", "nan", "null"}

    df = df.copy()
    df["numbered"] = pd.to_numeric(df.get("numbered", 0), errors="coerce").fillna(0).astype(int)

    if "bron_numId" not in df.columns:
        df["bron_numId"] = df.get("numId")

    # --- NIEUW: markeer uniforme blokken waar indented genegeerd moet worden ---
    # Uniform blok = contigu segment met:
    #   - opmaakprofiel in PROFILES
    #   - numbered == 1
    #   - overal zelfde opsommingstype (niet leeg)
    #   - overal zelfde bron_numId (geldig)
    df["__ignore_indent"] = 0
    n = len(df)
    i = 0
    while i < n:
        if (df.at[i, "opmaakprofiel"] in PROFILES) and (_safe_int(df.at[i, "numbered"], 0) == 1):
            prof = df.at[i, "opmaakprofiel"]
            j = i
            while j < n and (df.at[j, "opmaakprofiel"] == prof) and (_safe_int(df.at[j, "numbered"], 0) == 1):
                j += 1

            types = {_norm_type(df.at[k, "opsommingstype"]) for k in range(i, j)}
            brons = {str(df.at[k, "bron_numId"]) for k in range(i, j)}
            # Lege/ongeldige types of brons → niet uniform
            types_valid = (len(types) == 1) and (list(types)[0] != "")
            brons_valid = (len(brons) == 1) and _valid_bron(list(brons)[0])
            if types_valid and brons_valid:
                df.loc[i:j-1, "__ignore_indent"] = 1
            i = j
        else:
            i += 1
            
    # ---------- Case 1 ----------
    to_drop = []
    for i in range(1, len(df) - 1):
        row = df.iloc[i]
        if row.get("opmaakprofiel") in PROFILES:
            if not _valid_numid(row.get("numId")):
                prev = df.iloc[i - 1]
                nxt  = df.iloc[i + 1]
                if (prev.get("opmaakprofiel") in PROFILES and
                    nxt.get("opmaakprofiel")  in PROFILES and
                    _valid_numid(prev.get("numId")) and
                    str(prev.get("numId")) == str(nxt.get("numId"))):
                    df.at[i - 1, "tekst"] = str(prev["tekst"]).rstrip() + " " + str(row["tekst"]).lstrip()
                    logger.info('LET OP: Case 1 is wellicht nog niet goed.')
                    to_drop.append(i)
    if to_drop:
        df.drop(index=to_drop, inplace=True)
        df.reset_index(drop=True, inplace=True)

    # ---------- Case 2 ----------
    last_idx_by_profile       = {p: None for p in PROFILES}
    last_curr_numid_by_prof   = {p: None for p in PROFILES}
    last_bron_numid_by_prof   = {p: None for p in PROFILES}
    last_type_by_prof         = {p: None for p in PROFILES}

    for i in range(len(df)):
        row = df.iloc[i]
        p = row.get("opmaakprofiel")
        if p in PROFILES:
            has_curr = _valid_numid(row.get("numId"))
            if row.get("numbered", 0) == 1 and has_curr:
                
                cur_id   = row["numId"]
                cur_bid  = row.get("bron_numId")
                cur_type = _norm_type(row.get("opsommingstype"))

                if last_idx_by_profile[p] is None:
                    last_idx_by_profile[p]     = i
                    last_curr_numid_by_prof[p] = cur_id
                    last_bron_numid_by_prof[p] = cur_bid
                    last_type_by_prof[p]       = cur_type
                else:
                    prev_i   = last_idx_by_profile[p]
                    prev_id  = last_curr_numid_by_prof[p]
                    prev_bid = last_bron_numid_by_prof[p]
                    prev_type= last_type_by_prof[p]

                    false_new = (cur_type == prev_type) and (
                        str(cur_id) != str(prev_id) or str(cur_bid) != str(prev_bid)
                    )

                    if false_new:
                        j = i
                        while j < len(df):
                            same_profile = (df.at[j, "opmaakprofiel"] == p)
                            same_bad_bid = (str(df.at[j, "bron_numId"]) == str(cur_bid))
                            same_type    = (_norm_type(df.at[j, "opsommingstype"]) == cur_type)
                            j_valid      = _valid_numid(df.at[j, "numId"])
                            if same_profile and j_valid and same_bad_bid and same_type:
                                prev_level = _safe_int(df.at[j-1, "niveau"], 0)

                                # --- NIEUW: negeer indented in uniforme blokken ---
                                if _safe_int(df.at[j, "__ignore_indent"], 0) == 1:
                                    expected = prev_level  # geen verhoging/verlaging op basis van indent
                                else:
                                    prev_ind   = _safe_int(df.at[j-1, "indented"], 0)
                                    cur_ind    = _safe_int(df.at[j,   "indented"], 0)
                                    if cur_ind > prev_ind:
                                        expected = prev_level + 1
                                    elif cur_ind == prev_ind:
                                        expected = prev_level
                                    else:
                                        expected = max(0, prev_level - 1)

                                df.at[j, "numId"]  = prev_id
                                df.at[j, "niveau"] = expected
                                j += 1
                            else:
                                break

                    last_idx_by_profile[p]     = i
                    last_curr_numid_by_prof[p] = df.at[i, "numId"]
                    last_bron_numid_by_prof[p] = df.at[i, "bron_numId"]
                    last_type_by_prof[p]       = _norm_type(df.at[i, "opsommingstype"])
        else:
            for prof in PROFILES:
                last_idx_by_profile[prof]     = None
                last_curr_numid_by_prof[prof] = None
                last_bron_numid_by_prof[prof] = None
                last_type_by_prof[prof]       = None

    # ---------- Case 3 ----------
    last_anchor_idx = {p: None for p in PROFILES}
    to_drop = []
    for i, row in df.iterrows():
        p = row.get("opmaakprofiel") 
        if p in PROFILES:
            if (row.get("numbered", 0) == 1 and _valid_numid(row.get("numId"))) or row["fout"] == 2:
                last_anchor_idx[p] = i
            else:
                
                anchor = last_anchor_idx.get(p)
                if anchor is not None:
                    df.at[anchor, "fout"] = 2
                    df.at[anchor, "regel"] = 'Deze zin is samengevoegd met 1 of meerdere vervolgzinnen uit het brondocument.'
                    df.at[anchor, "tekst"] = str(df.at[anchor, "tekst"]).rstrip() + " " + str(row["tekst"]).lstrip()

                    textparts = ast.literal_eval(str(row["textparts"]))
                    textparts_anchor = ast.literal_eval(str(df.at[anchor, "textparts"]))
                    df.at[anchor, "textparts"] = textparts_anchor + textparts

                    number_of_textparts_anchor = len(textparts_anchor)
                    number_of_textparts = len(textparts)
                    
                    if number_of_textparts == 0:
                        number_of_textparts = 1
                        df.at[anchor, "textparts"].append(str(row["tekst"]).lstrip())
                        number_of_textparts_anchor = number_of_textparts_anchor + 1
                        
                    if df.at[anchor, "textpartformat_u"] == 'Ja':
                        textpartformats_u_anchor = ast.literal_eval(str(df.at[anchor, "textpartformats_u"]))
                    else:
                        textpartformats_u_anchor = ['None'] * number_of_textparts_anchor                        
                    if df.at[anchor, "textpartformat_b"] == 'Ja':
                        textpartformats_b_anchor = ast.literal_eval(str(df.at[anchor, "textpartformats_b"]))
                    else: 
                        textpartformats_b_anchor = ['None'] * number_of_textparts_anchor
                        
                    if df.at[anchor, "textpartformat_i"] == 'Ja':
                        textpartformats_i_anchor = ast.literal_eval(str(df.at[anchor, "textpartformats_i"]))
                    else:
                        textpartformats_i_anchor = ['None'] * number_of_textparts_anchor
                        

                    if row.get("textpartformat_u") == 'Ja':
                        df.at[anchor, "textpartformat_u"] = 'Ja'
                        textpartformats_u = row["textpartformats_u"]
                    else:
                        textpartformats_u = ['None'] * number_of_textparts
                
                    if row.get("textpartformat_b") == 'Ja':
                        df.at[anchor, "textpartformat_b"] = 'Ja'
                        textpartformats_b = ast.literal_eval(str(row["textpartformats_b"]))
                    else:
                        textpartformats_b = ['None'] * number_of_textparts
                        
                    if row.get("textpartformat_i") == 'Ja':
                        df.at[anchor, "textpartformat_i"] = 'Ja'
                        textpartformats_i = ast.literal_eval(str(row["textpartformats_i"]))
                    else:
                        textpartformats_i = ['None'] * number_of_textparts

                    if df.at[anchor, "textpartformat_u"] == 'Ja':
                        df.at[anchor, "textpartformats_u"] = textpartformats_u_anchor + textpartformats_u
                    if df.at[anchor, "textpartformat_b"] == 'Ja':
                        df.at[anchor, "textpartformats_b"] = textpartformats_b_anchor + textpartformats_b
                    if df.at[anchor, "textpartformat_i"] == 'Ja':
                        df.at[anchor, "textpartformats_i"] = textpartformats_i_anchor + textpartformats_i

                    to_drop.append(i)
        else:
            for prof in PROFILES:
                last_anchor_idx[prof] = None

    if to_drop:
        df.drop(index=sorted(set(to_drop)), inplace=True)
        df.reset_index(drop=True, inplace=True)

    # ---------- Case 4 ----------
    for i in range(1, len(df)):
        cur_prof  = df.at[i, "opmaakprofiel"] if "opmaakprofiel" in df.columns else None
        prev_prof = df.at[i-1, "opmaakprofiel"] if "opmaakprofiel" in df.columns else None
        if cur_prof in PROFILES and prev_prof == cur_prof:
            if _safe_int(df.at[i, "numbered"], 0) == 1 and _safe_int(df.at[i-1, "numbered"], 0) == 1:
                cur_type  = _norm_type(df.at[i,   "opsommingstype"])
                prev_type = _norm_type(df.at[i-1, "opsommingstype"])
                cur_bid   = str(df.at[i,   "bron_numId"])
                prev_bid  = str(df.at[i-1, "bron_numId"])
                if cur_type and cur_type == prev_type and cur_bid != prev_bid:
                    df.at[i, "niveau"] = _safe_int(df.at[i-1, "niveau"], 0)

    # ---------- Case 5 (hybride) ----------
    n = len(df)
    i = 1
    while i < n:
        cur_prof = df.at[i, "opmaakprofiel"] if "opmaakprofiel" in df.columns else None
        if cur_prof in PROFILES and _safe_int(df.at[i, "numbered"], 0) == 1:
            numid = df.at[i, "numId"]
            j = i
            while j < n and df.at[j, "opmaakprofiel"] == cur_prof and _safe_int(df.at[j, "numbered"], 0) == 1 and df.at[j, "numId"] == numid:
                j += 1

            # --- NIEUW: sla Case 5 over als dit segment volledig 'ignore_indent' is ---
            if bool(pd.to_numeric(df.loc[i:j-1, "__ignore_indent"], errors="coerce").fillna(0).astype(int).all()):
                i = j
                continue

            # Pak de niveaus in dit blok
            nivs = pd.to_numeric(df.loc[i:j-1, "niveau"], errors="coerce").fillna(0).astype(int)
            need_fix = (len(set(nivs)) == 1)

            for k in range(i, j):
                prev_ind = _safe_int(df.at[k-1, "indented"], 0)
                cur_ind  = _safe_int(df.at[k,   "indented"], 0)
                prev_lvl = _safe_int(df.at[k-1, "niveau"],   0)

                cur_type  = _norm_type(df.at[k,   "opsommingstype"])
                prev_type = _norm_type(df.at[k-1, "opsommingstype"])

                if cur_type and cur_type == prev_type:
                    if need_fix or (cur_ind > prev_ind and _safe_int(df.at[k, "niveau"], 0) == prev_lvl):
                        if cur_ind > prev_ind:
                            df.at[k, "niveau"] = prev_lvl + 1
                        elif cur_ind == prev_ind:
                            df.at[k, "niveau"] = prev_lvl
                        # bij daling: niet verlagen

            i = j
        else:
            i += 1

    # ---------- Case 6 ----------
    n = len(df)
    i = 0
    while i < n:
        if (df.at[i, "opmaakprofiel"] in PROFILES) and (_safe_int(df.at[i, "numbered"], 0) == 1):
            prof = df.at[i, "opmaakprofiel"]
            j = i
            while j < n and (df.at[j, "opmaakprofiel"] == prof) and (_safe_int(df.at[j, "numbered"], 0) == 1):
                j += 1

            lvls = pd.to_numeric(df.loc[i:j-1, "niveau"], errors="coerce").fillna(0).astype(int)
            inds = pd.to_numeric(df.loc[i:j-1, "indented"], errors="coerce").fillna(0).astype(int) if "indented" in df.columns else pd.Series([0]*(j-i))

            if len(inds) > 0 and inds.max() == inds.min() and lvls.nunique() == 1:
                base_level = _safe_int(df.at[i, "niveau"], 0)
                first_type = _norm_type(df.at[i, "opsommingstype"])
                first_bid  = str(df.at[i, "bron_numId"])
                if first_type and _valid_bron(first_bid):
                    path = [(first_type, first_bid)]
                    for k in range(i+1, j):
                        cur_type = _norm_type(df.at[k, "opsommingstype"])
                        cur_bid  = str(df.at[k, "bron_numId"])
                        if not cur_type or not _valid_bron(cur_bid):
                            continue

                        cur_pair = (cur_type, cur_bid)
                        prev_type, prev_bid = path[-1]

                        if cur_pair == (prev_type, prev_bid):
                            lvl = base_level + (len(path) - 1)
                        else:
                            try:
                                idx = next(idx for idx, pr in enumerate(path) if pr == cur_pair)
                                path = path[:idx+1]
                                lvl = base_level + idx
                            except StopIteration:
                                if cur_type != prev_type:
                                    path.append(cur_pair)
                                    lvl = base_level + (len(path) - 1)
                                else:
                                    path[-1] = cur_pair
                                    lvl = base_level + (len(path) - 1)

                        df.at[k, "niveau"] = int(lvl)
            i = j
        else:
            i += 1

    # Ruim hulpkolom op
    df.drop(columns=["__ignore_indent"], errors="ignore", inplace=True)

    return df



def correct_numbering(checked_document_df, input_document):
    # ---------- Helpers ----------
    def _to_int(x, default=0):
        """Robuust int: 'None', '', NaN -> default."""
        try:
            v = pd.to_numeric([x], errors='coerce')[0]
            if pd.isna(v):
                return default
            return int(v)
        except Exception:
            return default

    def _to_int_nullable(x):
        """Robuust int die None kan teruggeven (voor 'inputniveau' check)."""
        try:
            v = pd.to_numeric([x], errors='coerce')[0]
            if pd.isna(v):
                return None
            return int(v)
        except Exception:
            return None
    numbering_xml = 'word/numbering.xml'
    dh_bron = DocxHelper(input_document)
    numbering_xml_tree = dh_bron.get_xml(numbering_xml)
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # ---------- Normaliseer kolommen ----------
    # indented & numbered → int
    for col in ['indented', 'numbered']:
        if col in checked_document_df.columns:
            checked_document_df[col] = pd.to_numeric(checked_document_df[col], errors='coerce').fillna(0).astype(int)
        else:
            checked_document_df[col] = 0

    # niveau NIET casten naar int: bewaar NaN als "geen inputniveau"    
    if 'niveau' in checked_document_df.columns:
        checked_document_df['niveau'] = pd.to_numeric(checked_document_df['niveau'], errors='coerce')
    else:
        checked_document_df['niveau'] = pd.Series([pd.NA] * len(checked_document_df))

    if 'opsommingstype' not in checked_document_df.columns:
        checked_document_df['opsommingstype'] = None
        
    if 'bron_numId' not in checked_document_df.columns:
        checked_document_df['bron_numId'] = checked_document_df.get('numId')

    # 🎯 Snapshot van het oorspronkelijke input-niveau vóórdat we gaan schrijven
    niveau_input = checked_document_df['niveau'].copy()  # kan NaN bevatten
    

    lijst_teller = 0
    nieuwe_lijst = 'Ja'
    huidige_niveaus = []
    num_id_mapping = {}
    next_num_id = 1
    laatste_opsommingstype = None
    huidige_numId = None

    for index, row in checked_document_df.iterrows():
        # Sla speciale gevallen over
        if row['opmaakprofiel'] == 'OPOndertekening':
            continue
        elif (row['tekst'] in ['b e s l u i t:', 'b e s l u i t :', 'b   e   s   l   u   i   t   :']) and row['opmaakprofiel'] == 'OPAanhef':
            continue

        # Nieuwe lijst na titels
        if row['opmaakprofiel'] in ['OPArtikelTitel', 'OPParagraafTitel', 'OPHoofdstukTitel']:
            lijst_teller += 1
            nieuwe_lijst = 'Ja'
            huidige_niveaus = []
            laatste_opsommingstype = None

        if row['opmaakprofiel'] in ['OPArtikelTitel', 'OPParagraafTitel', 'OPHoofdstukTitel']:
            continue
            
            
        tekst = str(row.get("tekst", "")).replace('\xa0', ' ').lstrip()
        is_opsomming = _to_int(row.get('numbered', 0), 0) == 1
        opsommingstype = None

        # Herken handmatig getypte opsommingstekens
        if not is_opsomming:
            patterns = {
                'bullet': r"^(?:\-|\•|°)\s*",
                'kleine letter': r"^[a-z][\.\)]?\s+",
                'nummer_o': r"^\d+[oO°º][\.\)]?\s+",
                'nummer': r"^\d+[\.\)]?\s+",
                'hoofdletter': r"^[A-Z][\.\)]?\s+",
                'romeins cijfer klein': r"^(ix|iv|v?i{0,3})[\.\)]?\s+",
                'romeins cijfer groot': r"^(IX|IV|V?I{0,3})[\.\)]?\s+",
            }
            for type_key, pattern in patterns.items():
                if re.match(pattern, tekst):
                    zin_zonder_opsomming = re.sub(pattern, "", tekst)
                    checked_document_df.at[index, 'tekst'] = zin_zonder_opsomming
                    checked_document_df.at[index, 'numbered'] = 1
                    opsommingstype = type_key
                    is_opsomming = True
                    break

        if opsommingstype or is_opsomming:
            # Start/continue lijst
            if nieuwe_lijst == 'Ja':
                lijst_teller += 1
                nieuwe_lijst = 'Nee'
                huidige_numId = next_num_id
                next_num_id += 1
            else:
                huidige_numId = num_id_mapping.get(lijst_teller, next_num_id)

            if lijst_teller not in num_id_mapping:
                num_id_mapping[lijst_teller] = huidige_numId
            checked_document_df.at[index, 'numId'] = huidige_numId

            # Opsommingstype instellen (optioneel via numbering.xml)
            if opsommingstype:
                checked_document_df.at[index, 'opsommingstype'] = opsommingstype
            else:
                original_num_id = row.get('numId')
                abstract_num_id = None
                for num_elem in numbering_xml_tree.findall(f".//{{{namespace}}}num"):
                    if num_elem.get(f"{{{namespace}}}numId") == str(original_num_id):
                        abstract_num_elem = num_elem.find(f"{{{namespace}}}abstractNumId")
                        if abstract_num_elem is not None:
                            abstract_num_id = abstract_num_elem.get(f"{{{namespace}}}val")
                        break
                if abstract_num_id:
                    for abstract_num in numbering_xml_tree.findall(f".//{{{namespace}}}abstractNum"):
                        if abstract_num.get(f"{{{namespace}}}abstractNumId") == abstract_num_id:
                            ilvl = _to_int(row.get('niveau', 0), 0)
                            lvl_elem = abstract_num.find(f".//{{{namespace}}}lvl[@w:ilvl='{ilvl}']", namespaces={'w': namespace})
                            if lvl_elem is not None:
                                num_fmt = lvl_elem.find(f"{{{namespace}}}numFmt")
                                if num_fmt is not None:
                                    opsommingstype_value = num_fmt.get(f"{{{namespace}}}val")
                                    mapping = {
                                        "decimal": "nummer",
                                        "decimalZero": "nummer",
                                        "cardinalText": "nummer",
                                        "lowerLetter": "kleine letter",
                                        "upperLetter": "hoofdletter",
                                        "lowerRoman": "romeins cijfer klein",
                                        "upperRoman": "romeins cijfer groot",
                                        "bullet": "bullet",
                                    }
                                    opsommingstype = mapping.get(opsommingstype_value, opsommingstype_value)
                                    checked_document_df.at[index, 'opsommingstype'] = opsommingstype

            # --- Niveau bepaling (BELANGRIJK) ---
            # 0) Inputniveau-snapshot aanhouden
            
            input_lvl = _to_int_nullable(niveau_input.iat[index])            
            ind = _to_int(row.get('indented', 0), 0)
            if input_lvl is not None:
                # 1) Als inputniveau aanwezig is: respecteer dat
                niveau = input_lvl
                
#            if input_lvl is not None:
#                if ind > 0 and input_lvl == 0:
#                    # Correctie: als er indentatie is, maar niveau=0 in de input,
#                    # gebruik baseline in plaats van inputniveau
#                    niveau = ind
#                else:
#                    niveau = input_lvl                                
                
            else:
                # 2) Anders: baseline op regels/indented
                if row['opmaakprofiel'] == 'OPLid' and index > 0 and checked_document_df.at[index-1, 'opmaakprofiel'] == 'OPArtikelTitel':
                    niveau = 0
                else:
                    niveau = ind if ind > 0 else 0

                # 3) Stabilisatie: alleen als huidig inputniveau ontbreekt
                STABILIZE_PROFILES = {'OPLid', 'Lijstalinea'}
                if index > 0:
                    prev_prof = checked_document_df.at[index-1, 'opmaakprofiel']
                    cur_prof  = row['opmaakprofiel']
                    prev_ind  = _to_int(checked_document_df.at[index-1, 'indented'], 0)
                    cur_ind   = ind
                    if prev_prof == cur_prof and cur_prof in STABILIZE_PROFILES and prev_ind == cur_ind:
                        prev_level = _to_int(checked_document_df.at[index-1, 'niveau'], 0)
                        prev_input_lvl = _to_int_nullable(niveau_input.iat[index-1])
                        # Neem alleen over als vorige input óf ontbrak óf gelijk was aan berekende prev_level
                        if (prev_input_lvl is None) or (prev_input_lvl == prev_level):
                            niveau = prev_level

            checked_document_df.at[index, 'niveau'] = int(niveau)
            laatste_opsommingstype = opsommingstype
            
        else:
            huidige_niveaus = []
            nieuwe_lijst = 'Ja'

    corrected_document_df = checked_document_df    
    
    # Cleanup op het eind (jouw bestaande)
    corrected_document_df = post_numbering_cleanup(corrected_document_df)
    
    # Zorg dat niveau echt int is (geen 0.0)
    corrected_document_df['niveau'] = pd.to_numeric(corrected_document_df['niveau'], errors='coerce').fillna(0).astype(int)

    corrected_document_df.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_5.csv', sep=';', index=False)
    logger.info("Opsommingen in document gecontroleerd en gecorrigeerd.")
    return corrected_document_df




def change_Lijstalinea_na_OPLid(corrected_document_df):
    #Deze functie is toegevoegd omdat de LLMs getraind zijn op DROP output waar de opsommingen binnen 
    #leden als Lijstalinea zijn weergegeven. Dat moet steeds OPLid zijn.

    made_changes = False
    
    for i in range(len(corrected_document_df)):
        huidig_opmaakprofiel = corrected_document_df.loc[i, 'opmaakprofiel']
        
        if i > 0:
            vorig_opmaakprofiel = corrected_document_df.loc[i-1, 'opmaakprofiel']
        
            if huidig_opmaakprofiel == 'Lijstalinea' and vorig_opmaakprofiel == 'OPLid':
                corrected_document_df.loc[i, 'opmaakprofiel'] = 'OPLid'
                made_changes = True

            #Dit was voor een testje. Dit kan niet geatomatiseerd worden
#            if huidig_opmaakprofiel == 'StandaardAlinea' and vorig_opmaakprofiel == 'OPArtikelTitel' and corrected_document_df.loc[i, 'numbered'] == 1:
#                corrected_document_df.loc[i, 'opmaakprofiel'] = 'OPLid'
#                made_changes = True
                
    changed_opsommingen_df = corrected_document_df
    changed_opsommingen_df.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_6.csv', sep=';', index=False)
    
    if made_changes:
        logger.info(f"Lijstalinea's binnen OPLid aangepast naar OPLid.")
    
    return changed_opsommingen_df
    
    
    

def voeg_entries_toe_from_dataframe(numbering_xml_tree, read_new_numbering_df):
    # Haal de root van de XML op
    root = numbering_xml_tree.getroot()
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w15_namespace = "http://schemas.microsoft.com/office/word/2012/wordml"
    ET.register_namespace("w", namespace)
    ET.register_namespace("w15", w15_namespace)

    # Groepeer de data per numId
    grouped_df = read_new_numbering_df[read_new_numbering_df['numbered'] == 1].groupby('numId')

    # Houd een lijst bij voor de w:num elementen
    num_elements = []

    for num_id, group in grouped_df:
        # Converteer numId naar een string om te gebruiken als attribuutwaarde in XML
        abstract_num_id = str(num_id)

        # Maak een nieuw <w:abstractNum> element aan met w15:restartNumberingAfterBreak
        abstract_num = ET.Element(
            f"{{{namespace}}}abstractNum",
            attrib={
                f"{{{namespace}}}abstractNumId": abstract_num_id,
                f"{{{w15_namespace}}}restartNumberingAfterBreak": "0"
            }
        )
        ET.SubElement(abstract_num, f"{{{namespace}}}multiLevelType", attrib={f"{{{namespace}}}val": "hybridMultilevel"})

        # Voeg <w:lvl> elementen toe voor elk uniek niveau in de groep
        for _, row in group.drop_duplicates(subset=['niveau']).iterrows():
            niveau = row['niveau']
            
            # Bepaal de waarden voor numFmt en inspringing op basis van niveau en opsommingstype
            opsommingstype = row['opsommingstype']
            num_fmt_val = {
                "nummer": "decimal",
                "kleine letter": "lowerLetter",
                "bullet": "bullet",
                "hoofdletter": "upperLetter",
                "romeins cijfer klein": "lowerRoman",
                "romeins cijfer groot": "upperRoman",
                "nummer_o": "decimal"  # Gebruik decimal, maar met aangepast lvlText
            }.get(opsommingstype, "decimal")

            # Maak <w:lvl> element voor dit niveau
            lvl_elem = ET.SubElement(abstract_num, f"{{{namespace}}}lvl", attrib={f"{{{namespace}}}ilvl": str(niveau)})
            ET.SubElement(lvl_elem, f"{{{namespace}}}start", attrib={f"{{{namespace}}}val": "1"})  # Start altijd bij 1
            ET.SubElement(lvl_elem, f"{{{namespace}}}numFmt", attrib={f"{{{namespace}}}val": num_fmt_val})
            
            # Stel lvlText in afhankelijk van het opsommingstype
            if opsommingstype == "bullet":
                ET.SubElement(lvl_elem, f"{{{namespace}}}lvlText", attrib={f"{{{namespace}}}val": "-"})
            elif opsommingstype == "nummer_o":
                ET.SubElement(lvl_elem, f"{{{namespace}}}lvlText", attrib={f"{{{namespace}}}val": f"%{niveau + 1}°."})
            else:
                ET.SubElement(lvl_elem, f"{{{namespace}}}lvlText", attrib={f"{{{namespace}}}val": f"%{niveau + 1}."})

            # Voeg <w:ind> toe met aangepaste inspringing
            left_indent = 720 + (niveau * 720)
            pPr = ET.SubElement(lvl_elem, f"{{{namespace}}}pPr")
            ET.SubElement(pPr, f"{{{namespace}}}ind", attrib={f"{{{namespace}}}left": str(left_indent), f"{{{namespace}}}hanging": "360"})

        # Voeg het <w:abstractNum> element toe aan de root
        root.append(abstract_num)

        # Voor de <w:num> elementen, maak een element en voeg het toe aan de num_elements lijst
        num = ET.Element(f"{{{namespace}}}num", attrib={f"{{{namespace}}}numId": abstract_num_id})
        ET.SubElement(num, f"{{{namespace}}}abstractNumId", attrib={f"{{{namespace}}}val": abstract_num_id})
        
        # Voeg <w:lvlOverride> toe voor herstarten bij elk nieuw niveau
        lvl_override = ET.SubElement(num, f"{{{namespace}}}lvlOverride", attrib={f"{{{namespace}}}ilvl": "0"})
        ET.SubElement(lvl_override, f"{{{namespace}}}startOverride", attrib={f"{{{namespace}}}val": "1"})

        # Voeg <w:num> toe aan de lijst om later aan het einde te plaatsen
        num_elements.append(num)

    # Voeg alle <w:num> elementen aan het einde van de root toe
    for num_element in num_elements:
        root.append(num_element)

    # Sla het gewijzigde XML-document op of retourneer het
    return numbering_xml_tree

def escape_placeholders(text: str) -> str:
    """
    Escape <...> placeholders in gewone tekst zodat ze geldig XML worden.
    Voorbeeld: 'De inrichting <naam inrichting>' -> 'De inrichting &lt;naam inrichting&gt;'
    """
    return text.replace("<", "&lt;").replace(">", "&gt;")


def generate_xml(checked_document_df):

    if os.path.exists("temp_files/temp.xml"):
        os.remove("temp_files/temp.xml")

    if os.path.exists("temp_files/tempcomments.xml"):
        os.remove("temp_files/tempcomments.xml")
    
    tempxml = open("temp_files/temp.xml", "x")
    tempcommentsxml = open("temp_files/tempcomments.xml", "x")
    foutmelding = ''
    comments_generated = False
    fout_teller = -1
    comments_include_errors = False
    
    
    for index, row in checked_document_df.iterrows():
        
        wrPr_start = '<w:rPr>'
        wrPr_end = '</w:rPr>'
        bold_tag = ''
        italic_tag = ''
        underline_tag = ''
        
        #check of er fouten geconstateerd waren
        if row['fout'] == 1 or row['fout'] == 2:
            if row['fout'] == 1:
                comments_include_errors = True
            comments_generated = True                
            fout_teller = fout_teller + 1
            foutmelding = row['regel']
            highlight_text_1 = '<w:highlight w:val="yellow"/>'
            highlight_text_2 = '<w:highlight w:val="yellow"/>'
            highlight_text_3 = '<w:rPr><w:highlight w:val="yellow"/></w:rPr>'
            comment_start = f'<w:commentRangeStart w:id="{fout_teller}"/>'
            comment_end = f'<w:commentRangeEnd w:id="{fout_teller}"/><w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:commentReference w:id="{fout_teller}"/></w:r>'

            comment = f'<w:comment w:id="{fout_teller}" w:author="KOOP Converter" w:initials="KC"><w:p><w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r><w:r><w:rPr><w:color w:val="000000"/></w:rPr><w:t>{foutmelding}</w:t></w:r></w:p></w:comment>'
            tempcommentsxml.write(comment) 
        else:
            foutmelding = ''
            highlight_text_1 = ''
            highlight_text_2 = ''
            highlight_text_3 = ''
            comment_start = ''
            comment_end = ''
            
            
        tempxml.write(f'<w:p><w:pPr><w:pStyle w:val="{row["opmaakprofiel"]}"/>')

        if row['numbered'] == 1:
            tempxml.write(f'<w:numPr><w:ilvl w:val="{row["niveau"]}" />')
            tempxml.write(f'<w:numId w:val="{row["numId"]}"/></w:numPr>')

        tempxml.write(f'</w:pPr>{comment_start}')        
                
        if (row['textpartformat_u'] == 'Nee' and row['textpartformat_b'] == 'Nee' and row['textpartformat_i'] == 'Nee') or row["opmaakprofiel"] == 'OPArtikelTitel':
            #gebruik hele paragraph            
            tekst_lb = escape_placeholders(row["tekst"]).replace("\n", '</w:t><w:br/><w:t xml:space="preserve">')
            text_cleaned = re.sub(r' {2,}', ' ', tekst_lb) #spaties in midden weghalen
            if row["bold"] == 1:
                bold_tag = '<w:b/>'
            if row["italic"] == 1:
                italic_tag = '<w:i/>'
            if row["underlined"] == 1:
                underline_tag = '<w:u w:val="single"/>'  
                
            tempxml.write(f'<w:r>{wrPr_start}{highlight_text_1}{bold_tag}{italic_tag}{underline_tag}{wrPr_end}<w:t xml:space="preserve">{text_cleaned}</w:t></w:r>')
        else:
            #gebruik text parts en opmaak per part
            for i in range(len(row['textparts'])):
                opmaak = 0

                #alleen als er voor de huidige text part opmaak is. 
                if (row['textpartformat_b'] == 'Ja' and row['textpartformats_b'][i] != 'None' and row['textpartformats_b'][i] != False) or (row['textpartformat_u'] == 'Ja' and row['textpartformats_u'][i] != 'None' and row['textpartformats_u'][i] != False) or (row['textpartformat_i'] == 'Ja' and row['textpartformats_i'][i] != 'None' and row['textpartformats_i'][i] != False):
                    tempxml.write(f'<w:r><w:rPr>')
                if row['textpartformat_b'] == 'Ja' and row['textpartformats_b'][i] != 'None' and row['textpartformats_b'][i] != False:
                    opmaak = 1
                    tempxml.write(f'<w:b/><w:bCs/>')
                if row['textpartformat_i'] == 'Ja' and row['textpartformats_i'][i] != 'None' and row['textpartformats_i'][i] != False:
                    opmaak = 1
                    tempxml.write(f'<w:i/><w:iCs/>')
                if row['textpartformat_u'] == 'Ja' and row['textpartformats_u'][i] != 'None' and row['textpartformats_u'][i] != False:
                    opmaak = 1
                    if int(row['textpartformats_u'][i]) == 1:
                        tempxml.write(f'<w:u w:val="single" />')
                    elif int(row['textpartformats_u'][i]) == 3:
                        tempxml.write(f'<w:u w:val="double" />')

                tekst_met_linebreak = escape_placeholders(row["textparts"][i]).replace("\n", '</w:t><w:br/><w:t xml:space="preserve">')
                text_cleaned_met_lb = re.sub(r' {2,}', ' ', tekst_met_linebreak) #spaties in midden weghalen
                        
                if opmaak == 1:
                    tempxml.write(f'{highlight_text_2}</w:rPr><w:t xml:space="preserve">{text_cleaned_met_lb}</w:t></w:r>')
                else:
                    tempxml.write(f'<w:r>{highlight_text_3}<w:t xml:space="preserve">{text_cleaned_met_lb}</w:t></w:r>')

        tempxml.write(f'{comment_end}</w:p>')

        #Na elke OPAanhef regel moet een lege regel
        if row["opmaakprofiel"] == 'OPAanhef':
            tempxml.write(f'<w:p><w:pPr><w:pStyle w:val="OPAanhef" /></w:pPr></w:p>')
            


    tempxml.close()         
    file_header = open("template/xml/document_header.xml", "r")
    file_body = open("temp_files/temp.xml", "r")
    file_footer = open("template/xml/document_footer.xml", "r")

    header = file_header.read()
    body = file_body.read()
    footer = file_footer.read()

    file_header.close()
    file_body.close()
    file_footer.close()

    document_xml = open("resultaat/xml/document.xml", "w")
    document_xml.write(header + body + footer)
    document_xml.close()       
    os.remove("temp_files/temp.xml")
    
    tempcommentsxml.write('</w:comments>')
    tempcommentsxml.close()
    file_comments_header = open("template/xml/comments_header.xml", "r")
    file_comments_body = open("temp_files/tempcomments.xml", "r")
    
    comments_header = file_comments_header.read()
    comments_body = file_comments_body.read()
    
    file_comments_header.close()
    file_comments_body.close()
    
    comments_xml = open("resultaat/xml/comments.xml", "w")
    comments_xml.write(comments_header + comments_body)
    comments_xml.close()       
    os.remove("temp_files/tempcomments.xml")
    
    logger.info(f"XML-inhoud van het document gegenereerd.")

    return comments_generated, comments_include_errors

    
def generate_word(comments_generated, input_document, output_document):
    #dit is het lege DROP template
    input_docx = 'template/OP_Stijl Compleet Besluit v2.5_0.docx'
    #dit is de naam van de xml die vervangen gaat worden. Als je een docx unzipped dan worden er diverse
    #sub-directories aangemaakt, waaronder de directory 'word' waar dat document.xml in staat. 
    filename = 'word/document.xml'
    numbering_xml = 'word/numbering.xml'
    numbering_tmp = 'template/xml/numbering.xml'
    
    #dit is jouw xml. Dus hier moet je de directory 'resultaat' even vervangen door hoe jij het genoemd hebt.
    new_xml = 'resultaat/xml/document.xml'
    
    # Lees originele XML in als tekst
    with open(new_xml, "r", encoding="utf-8") as f:
        xml_text = f.read()

    # Vervang losse & (geen deel van geldige XML-entiteit) door &amp;
    xml_text_fixed = re.sub(r'&(?!amp;|lt;|gt;|quot;|apos;)', '&amp;', xml_text)
    
    # Overschrijf bestand met de gefixte inhoud
    with open(new_xml, "w", encoding="utf-8") as f:
        f.write(xml_text_fixed)    
        
    #hier creeer je een object. Dit object geef je zelf een naam, zoals een variable, in dit geval 'dh'
    #het is een object van de klasse 'DocxHelper' die hierboven gemaakt is. 
    #de enige input die deze klasse vraagt om een object te maken is een bestandsnaam van het input docx bestand,
    #in dit geval dus het lege DROP sjabloon
    dh = DocxHelper(input_docx)
    
    
    numbering_xml_tree = ET.parse(numbering_tmp)
    
    read_new_numbering_df = pd.read_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_5.csv', sep=';')

    
    
    # Converteer de kolommen 'numbered' en 'numId' naar integers, vervang 'None' en niet-numerieke waarden door 0
    read_new_numbering_df['numbered'] = pd.to_numeric(read_new_numbering_df['numbered'], errors='coerce').fillna(0).astype(int)
    read_new_numbering_df['numId'] = pd.to_numeric(read_new_numbering_df['numId'], errors='coerce').fillna(0).astype(int)
    read_new_numbering_df['niveau'] = pd.to_numeric(read_new_numbering_df['niveau'], errors='coerce').fillna(0).astype(int)

    
    numbering_xml_tree_ext = voeg_entries_toe_from_dataframe(numbering_xml_tree, read_new_numbering_df)
    

    #hier creeer je een object van het type xml-tree. Dat doe je met de functie parse van de klasse ET, zie 
    #derde import statement. Dit heb je nodig omdat de DocxHelper class een xml-tree verwacht bij het vervangen
    #van een xml document.
    systeminfo = platform.uname()
    if systeminfo[0] == 'Windows':
        with open(new_xml, encoding="windows-1252") as filexxx:
            xml_tree = ET.parse(filexxx)
    else:
        xml_tree = ET.parse(new_xml)
    
    #hier roep je de functie set_xml aan van jouw dh object. Deze functie verwacht drie parameters
    #1. de naam van het bestand dat je wilt vervangen. In dit geval dus 'word/document.xml'. Let op dit is dus het
    #   oude xml bestand uit de template
    #2. de nieuwe xml als xml-tree. Dat is dus het object xml_tree dat je net hebt aangemaakt op basis van jouw
    #.  document.xml bestand
    #3. de naam van het bestand dat hij moet aanmaken. In dit geval dus outputdoctest.docx.
    #dh.set_xml(filename, xml_tree, output_docx)

    if comments_generated:
        new_comments_xml = 'resultaat/xml/comments.xml'
        dh.set_xmls_comments(filename, new_comments_xml, numbering_xml, xml_tree, numbering_xml_tree_ext, output_document)
    else:
        dh.set_xmls(filename, numbering_xml, xml_tree, numbering_xml_tree_ext, output_document)
        
        

def get_paragraphs(document):
    
    docx_dict = {}
    i = 0
    i_2 = 0
    p_index = 0
    
    tables_info = []  # Hier slaan we de tabelinformatie op
    images_info = []  # Hier slaan we de info over afbeeldingen op

    # Loop door de elementen in volgorde in het document
    for element in document.element.body:
        if element.tag.endswith("tbl"):  # Detecteert een tabel
            i_2 += 1
            tables_info.append({'volgnummer': i_2, 'tekst': "TABEL GEVONDEN."})

        elif element.tag.endswith("p"):  # Detecteert een paragraaf
            paragraph = document.paragraphs[p_index]

            # Check of er een afbeelding in deze paragraaf zit
            if paragraph._element.xpath(".//*[local-name()='drawing']") or paragraph._element.xpath(".//*[local-name()='pict']"):
                # Zoek naar de naam van de afbeelding (docPr/@name)
                docPr_nodes = paragraph._element.xpath(".//*[local-name()='docPr']")
                img_name = docPr_nodes[0].get("name") if docPr_nodes else None

                # Zoek naar de relatie-id (blip/@r:embed)
                blip_nodes = paragraph._element.xpath(".//*[local-name()='blip']")
                rel_id = None
                if blip_nodes:
                    # De attribuutnaam bevat namespace, dus zo ophalen:
                    for attr in blip_nodes[0].attrib:
                        if attr.endswith("embed"):
                            rel_id = blip_nodes[0].attrib[attr]

                i_2 += 1
                images_info.append({
                    'volgnummer': i_2,
                    'tekst': "AFBEELDING GEVONDEN.",
                    'naam': img_name,
                    'rel_id': rel_id
                })

            test_paragraph = paragraph.text.strip(' ')
            if (len(test_paragraph) > 0) and not test_paragraph.isspace():
                i_2 += 1

            p_index += 1

    
    for paragraph in document.paragraphs:
        test_paragraph = paragraph.text.strip(' ')
        if (len(test_paragraph) > 0) and not test_paragraph.isspace():
            i = i + 1
            paragraph_sentence = {}
            paragraph_sentence['volgnummer'] = i
            paragraph_sentence['tekst'] = paragraph.text  
                        
            if paragraph.paragraph_format.left_indent == None:
                paragraph_sentence['indented'] = 0
                paragraph_sentence['indentation'] = 0
            else:
                if paragraph.paragraph_format.left_indent > 0:
                    paragraph_sentence['indented'] = 1
                    paragraph_sentence['indentation'] = paragraph.paragraph_format.left_indent
                else:
                    paragraph_sentence['indented'] = 0
                    paragraph_sentence['indentation'] = 0
                    
            try:
                if paragraph.runs[0].font.bold == True:
                    paragraph_sentence['bold'] = 1
                else:
                    paragraph_sentence['bold'] = 0
            except:
                paragraph_sentence['bold'] = 0
                
            try:
                if paragraph.runs[0].font.italic == True:
                    paragraph_sentence['italic'] = 1
                else:
                    paragraph_sentence['italic'] = 0
            except:
                paragraph_sentence['italic'] = 0
                    
            try:
                if paragraph.runs[0].font.underline == True:
                    paragraph_sentence['underlined'] = 1
                else:
                    paragraph_sentence['underlined'] = 0
            except:
                paragraph_sentence['underlined'] = 0
                
            try:
                paragraphformat = paragraph._element.pPr.numPr
            except:
                paragraphformat = None
                
            if (paragraphformat == None):
                paragraph_sentence['genummerd'] = 0
                paragraph_sentence['niveau'] = 'None'
                paragraph_sentence['numId'] = 'None'                
            else:
                paragraph_sentence['genummerd'] = 1
                paragraph_sentence['niveau'] = paragraphformat.ilvl.val
                paragraph_sentence['numId'] = paragraphformat.numId.val                     
        
            #Er kan ook een deel van een paragraph opmaak hebben.
            #Dan zijn er meerdere runs.
            textparts = []
            textpartformats_u = []
            textpartformats_b = []
            textpartformats_i = []
            textpartformat_u = 'Nee'
            textpartformat_b = 'Nee'
            textpartformat_i = 'Nee'            
            if len(paragraph.runs) > 1:
                #iterate over the parts
                for textpart in paragraph.runs:
                    #put textpart on the list
                    textparts.append(textpart._r.text)
                    
                    # try for underline
                    try: 
                        test = textpart._r.rPr.u.val
                        textpartformat_u = 'Ja'                        
                        textpartformats_u.append(test.value)
                    except:
                        textpartformats_u.append('None')

                    # try for bold
                    try: 
                        test = textpart._r.rPr.b.val
                        textpartformat_b = 'Ja'
                        textpartformats_b.append(textpart._r.rPr.b.val)    
                    except:
                        textpartformats_b.append('None')
                        
                    # try for italic
                    try: 
                        test = textpart._r.rPr.i.val
                        textpartformat_i = 'Ja'
                        textpartformats_i.append(textpart._r.rPr.i.val)    
                    except:
                        textpartformats_i.append('None')

            paragraph_sentence['textparts'] = textparts
            paragraph_sentence['partformat_u'] = textpartformat_u
            paragraph_sentence['partformat_b'] = textpartformat_b
            paragraph_sentence['partformat_i'] = textpartformat_i
            

            if textpartformat_u == 'Ja':
                paragraph_sentence['partformats_u'] = textpartformats_u
            if textpartformat_b == 'Ja':
                paragraph_sentence['partformats_b'] = textpartformats_b
            if textpartformat_i == 'Ja':
                paragraph_sentence['partformats_i'] = textpartformats_i
            
            docx_dict[i] = paragraph_sentence
    return docx_dict, tables_info, images_info

def combine_predictions_chatgpt(prediction_batches):
    # Totaal aantal zinnen afgeleid van het aantal voorspelling batches en hun lengte
    total_sentences = len(prediction_batches) + len(prediction_batches[0]) - 1
    
    # Maak een lege dictionary om de voorspellingen per zin op te slaan
    sentence_predictions = {i: [] for i in range(1, total_sentences + 1)}

    # Loop door elke batch van voorspellingen
    for i, batch in enumerate(prediction_batches):
        # Voor elk element in de batch, voeg de voorspelling toe aan de juiste zin
        for j, prediction in enumerate(batch):
            sentence_index = i + j + 1
            if sentence_index <= total_sentences:
                sentence_predictions[sentence_index].append(prediction)

    return sentence_predictions


def get_model_prediction(message):
    OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]
    client2 = OpenAI(api_key=OPENAI_API_KEY)

    
    completion2 = client2.chat.completions.create(
        model="ft:gpt-4o-mini-2024-07-18:spynk-consulting:dk-opmaak-met-features:AG81RAxG", 
        messages=message
    )
    return completion2.choices[0].message.content

def calculate_final_predictions(all_predictions):
    # Gewichtsschema voor verschillende aantallen voorspellingen
    weight_schemes = {
        5: [0.1, 0.1, 0.6, 0.1, 0.1],
        4: [0.2, 0.5, 0.2, 0.1],  
        3: [0.3, 0.4, 0.3],
        2: [0.6, 0.4],
        1: [1.0]
    }

    # Functie om de gewogen voorspelling te berekenen
    def calculate_weighted_prediction(predictions, weight_scheme):
        # Gebruik een default dictionary om de gewogen stemmen te tellen
        weighted_counts = defaultdict(float)
        
        # Voor elke voorspelling, pas het juiste gewicht toe
        for i, prediction in enumerate(predictions):
            weight = weight_scheme[i]  # Haal het gewicht op
            weighted_counts[prediction] += weight  # Voeg het gewicht toe voor dit label
        
        # Kies het label met de hoogste gewogen score
        final_prediction = max(weighted_counts, key=weighted_counts.get)
        return final_prediction

    # Dictionary om de uiteindelijke voorspellingen per zin op te slaan
    final_predictions = {}

    # Verwerk elke zin en zijn voorspellingen
    for sentence_id, predictions in all_predictions.items():
        num_predictions = len(predictions)
        
        # Haal het juiste gewichtsschema op voor dit aantal voorspellingen
        weight_scheme = weight_schemes[num_predictions]
        
        # Bereken de gewogen voorspelling
        final_prediction = calculate_weighted_prediction(predictions, weight_scheme)
        
        # Sla het resultaat op in de uiteindelijke voorspellingen (geen lijst meer)
        final_predictions[sentence_id] = final_prediction

    return final_predictions


def verwerk_bestand(input_document, output_dir, modelkeuze):
    
    global datum_map
    base_dir = Path("resultaat/csv")
    datum_map = datetime.now().strftime("%Y%m%d")
    csv_dir = base_dir / datum_map
    csv_dir.mkdir(parents=True, exist_ok=True)
    
    #check if document is .doc
    file_path = Path(input_document)
    global bestandsnaam_zonder_ext 
    bestandsnaam_zonder_ext = file_path.stem
    if file_path.suffix == ".doc":
        input_document_docx = str(file_path.with_suffix(".docx"))
        convert(input_document, input_document_docx)
        os.remove(input_document)
        input_document = input_document_docx
        logger.info(f"Document is van .doc naar .docx geconverteerd.")
    
    try:
        
        #Dit is nodig ongeacht het model. Hij maakt hier een docx object aan en stopt daarna alle zinnen in
        #de variabele document_sentences.
        document = Document(input_document)
        document_sentences, tabellen, afbeeldingen = get_paragraphs(document)
        # Controleer of document_sentences een dictionary is en zet het om naar een lijst
        if isinstance(document_sentences, dict):
            document_sentences = list(document_sentences.values())
        # Maak een dictionary om voorspellingen voor elke zin op te slaan
        predictions_per_sentence = {i: [] for i in range(1, len(document_sentences) + 1)}

        # Itereer over het document met een sliding window van 5 zinnen
        window_size = 5
        num_sentences = len(document_sentences)
        total_iterations = num_sentences - window_size + 1
        
        logger.info(f"Document ingelezen: {len(document_sentences)} zinnen zijn verwerkt.")
        
        tables_added = False
        if len(tabellen) > 0:
            logger.info(f'Aantal gevonden tabellen: {len(tabellen)}')
            tables_added = True

        images_added = False
        if len(afbeeldingen) > 0:
            logger.info(f'Aantal gevonden afbeeldingen: {len(afbeeldingen)}')
            images_added = True
            
        
        if modelkeuze == "BERT":
            
            #initialize BERT, returns model on mps in eval
            model, tokenizer, device = init_model_BERT()            
            logger.info(f"AI Model geïnitialiseerd: {modelkeuze}")
            
            for start_idx in range(total_iterations):


                # Pak 5 opeenvolgende zinnen
                window_sentences = document_sentences[start_idx:start_idx + window_size]
                # Combineer de zinnen met het [SEP] token
                combined_text = f" {tokenizer.sep_token} ".join([sentence['tekst'] for sentence in window_sentences])

                # Tokenize de gecombineerde tekst
                tokenized_input = tokenizer(combined_text, padding='max_length', truncation=True, max_length=512, return_tensors='pt')

                # Extract features zoals eerder gedaan
                bold_features = torch.tensor([[sentence['bold'] for sentence in window_sentences]], dtype=torch.float).to(device)
                indented_features = torch.tensor([[sentence['indented'] for sentence in window_sentences]], dtype=torch.float).to(device)
                italic_features = torch.tensor([[sentence['italic'] for sentence in window_sentences]], dtype=torch.float).to(device)
                underlined_features = torch.tensor([[sentence['underlined'] for sentence in window_sentences]], dtype=torch.float).to(device)
                numbered_features = torch.tensor([[sentence['genummerd'] for sentence in window_sentences]], dtype=torch.float).to(device)
                volgnummer_features = torch.tensor([[sentence['volgnummer'] for sentence in window_sentences]], dtype=torch.float).to(device)

                # Voer de input door het model
                input_ids = tokenized_input['input_ids'].to(device)
                attention_mask = tokenized_input['attention_mask'].to(device)

                with torch.no_grad():
                    outputs = model(input_ids=input_ids, attention_mask=attention_mask,
                                    bold_features=bold_features, indented_features=indented_features,
                                    italic_features=italic_features, underlined_features=underlined_features,
                                    numbered_features=numbered_features, volgnummer_features=volgnummer_features)

                # Verzamel voorspellingen per zin in het window
                for i, logit in enumerate(outputs[0]):  # itereren over de zinnen in het window
                    probabilities = torch.softmax(logit, dim=-1).cpu().numpy()  # Zet logits om in waarschijnlijkheden

                    # Index van de zin in het volledige document
                    sentence_index = start_idx + i + 1  

                    # Voeg toe: tuple van de waarschijnlijkheden en de positie in het sliding window
                    predictions_per_sentence[sentence_index].append((probabilities, i))
                    
            final_predictions, all_predictions = get_final_predictions(predictions_per_sentence)
            
        if modelkeuze == "ChatGPT":
            
            logger.info(f"AI Model niet geïnitialiseerd, model niet meer ondersteund: {modelkeuze}")
            
        
        logger.info(f"Opmaakprofielen per zin bepaald.")        
        predicted_document_df = predict_document(final_predictions, document_sentences)
        
        checked_document_df = check_predicted_document(predicted_document_df)
        annotated_document_df = check_structuur(checked_document_df, tabellen, afbeeldingen)
                
        #correctie voorspellingen
        chatgpt_corrected = apply_corrections_by_windows(annotated_document_df, corrector=chatgpt_corrector)        
        chatgpt_corrected.to_csv(f'resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_4a.csv', sep=';', index=False)

    
        # Her-check structuur:
        df_rechecked = check_structuur_recheck(chatgpt_corrected, write_csv_path=f"resultaat/csv/{datum_map}/{bestandsnaam_zonder_ext}_4b.csv")
        corrected_document_df = correct_numbering(df_rechecked, input_document)
        changed_opsommingen_df = change_Lijstalinea_na_OPLid(corrected_document_df)
        comments_generated, comments_include_errors = generate_xml(changed_opsommingen_df)
        

        inp = Path(input_document)
        filename = inp.name   
        output_document = f"inlaad_{filename}"            
        output_document_file = os.path.join(output_dir, output_document)
        generate_word(comments_generated, input_document, output_document_file)        
        
        
        logger.info(f"Word document gegenereerd op basis van KOOP template.")
        
        if comments_include_errors:
            logger.info(f"LET OP! Zie comments in het document voor fouten in de opmaakprofielen.")
        if tables_added:
            logger.info(f"LET OP! Zie comments in verband met toe te voegen tabellen.")
        if images_added:
            logger.info(f"LET OP! Zie comments in verband met toe te voegen afbeeldingen.")

        
    except Exception as e:
        logger.info(f"Fout bij het verwerken van bestand: {str(e)}")
